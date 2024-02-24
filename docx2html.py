import subprocess
import os
import re
import json
from bs4 import BeautifulSoup
import ratta_functions
from docx import Document
import glob
import shutil
import json
import os
import yaml
import re
import csv
import random
# import g4f
import docx2txt
import pypandoc
import docx
from itertools import combinations
import shutil
import glob
from graphviz import Digraph
import xml.etree.ElementTree as ET
import logging
from intelligence import Brain
import time
from indicparser import graphemeParser

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Inches
import json
import pandas as pd

import os
import logging
from pathlib import Path

def update_font_style_in_docx(input_file_path, output_file_path):

    def set_column_widths_auto(doc):
        for table in doc.tables:
            for column in table.columns:
                for cell in column.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.get_or_add_tcW()
                    tcW.type = 'auto'
                    tcW.w = 0

    def set_column_widths_manually(doc):
        def get_max_column_width(table, col_idx):
            max_width = 0
            for row in table.rows:
                cell = row.cells[col_idx]
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        max_width = max(max_width, len(
                            graphemeParser("hindi").process(run.text)))

            return max_width

        for table in doc.tables:
            table.autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            num_columns = len(table.columns)
            for col_idx in range(num_columns):
                max_width = get_max_column_width(table, col_idx)
                if max_width > 30:
                    max_width = 30
                # Convert character count to inches (adjust as needed)
                max_width_inches = max_width * 0.12  # Adjust this factor as needed
                # writeLog(max_width)
                table.columns[col_idx].width = Inches(max_width_inches)

    def set_cell_border(doc):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.get_or_add_tcW()
                    tcW.type = 'auto'
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is None:
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                    for border in ['top', 'left', 'bottom', 'right']:
                        tcBorder = OxmlElement(f'w:{border}')
                        tcBorder.set(qn('w:val'), 'single')
                        tcBorder.set(qn('w:sz'), '4')
                        tcBorder.set(qn('w:space'), '0')
                        tcBorder.set(qn('w:color'), '000000')
                        tcBorders.append(tcBorder)

    def set_font(paragraph, font_name, font_size):
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            # Ensure the font name is set correctly
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)
            rFonts.set(qn('w:eastAsia'), font_name)
            rFonts.set(qn('w:cs'), font_name)
            rPr.append(rFonts)

    # Load the document
    doc = Document(input_file_path)
    font_name = 'Sahitya'
    font_size = 12

    # Change the font for all paragraphs in the document
    for paragraph in doc.paragraphs:
        set_font(paragraph, font_name, font_size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_font(paragraph, font_name, font_size)

    # Format all tables in the document

    set_cell_border(doc)
    set_column_widths_auto(doc)
    # set_column_widths_manually(doc)

    # Save the modified document
    doc.save(output_file_path)

    return 0

def backend_clean_html(html):
    def backend_replace_options_in_table(match):
        table_html = match.group(0)
        # More specific replacement within <td> or <th> tags only
        table_html = re.sub(r'\(a\)', '(A)', table_html,
                            flags=re.IGNORECASE)
        table_html = re.sub(r'\(b\)', '(B)', table_html,
                            flags=re.IGNORECASE)
        table_html = re.sub(r'\(c\)', '(C)', table_html,
                            flags=re.IGNORECASE)
        table_html = re.sub(r'\(d\)', '(D)', table_html,
                            flags=re.IGNORECASE)
        return table_html

    # Apply the function to each table in the HTML
    html = re.sub(r'(<table.*?>.*?</table>)',
                    backend_replace_options_in_table, html, flags=re.DOTALL)

    html = re.sub(r'<((?:strong|sup|sub|mark|em|u|del))><br />\s*\n</\1>',
                    '<br />\n', html, flags=re.MULTILINE)
    html = re.sub(r'<((?:strong|sup|sub|mark|em|u|del))><br />\s*\n</\1>',
                    '<br />\n', html, flags=re.MULTILINE)

    # Continue with the existing cleanup operations
    html = re.sub(r'^(\d+\.\))', r'</p>\n<p>\1', html, flags=re.MULTILINE)
    html = re.sub(r'^\((a|b|c|d)\)', r'</p>\n<p>(\1)',
                    html, flags=re.MULTILINE)
    html = re.sub(r'^Ans\.', r'</p>\n<p>Ans.', html, flags=re.MULTILINE)
    html = re.sub(r'^Exp:', "</p>\n<p>Exp:", html, flags=re.MULTILINE)

    html = re.sub(r'(?<!</p>)\n(?!<p>)', ' ', html)
    html = re.sub(r'</p>\n<p>', "\n", html, flags=re.DOTALL)

    html = re.sub(r'<br />', "\n", html, flags=re.DOTALL)
    html = re.sub(r'<strong>\s*\n\s*</strong>', "", html, flags=re.DOTALL)
    html = re.sub(r'<u>\s*\n\s*</u>', "", html, flags=re.DOTALL)
    html = re.sub(r'<em>\s*\n\s*</em>', "", html, flags=re.DOTALL)
    html = re.sub(r'<del>\s*\n\s*</del>', "", html, flags=re.DOTALL)
    html = re.sub(r'<mark>\s*\n\s*</mark>', "", html, flags=re.DOTALL)
    # Replace HTML entities with spaces
    html = re.sub(r'&[a-z]+;', ' ', html)
    html = re.sub(r'\n<math', ' <math', html)
    html = re.sub(r'\n</math>', ' </math>\n', html)
    html = re.sub(r'</math><br/>', ' </math> ', html)
    html = re.sub(r'<body>.<p>', '<body>\n', html, flags=re.DOTALL)
    # Replace multiple newlines with a single newline
    html = re.sub(r'\n\n+', '\n', html)
    html = re.sub(r'<p>', '', html)
    html = re.sub(r'</p>', '', html)
    html = re.sub(r'</body> </html>', '', html)
    return html

def backend_extract_html_elements(text):

    def parse_content(contents):
        for content in contents:
            if isinstance(content, str) or content.name == 'math' or content.name == 'strong' or content.name == 'em' or content.name == 'u' or content.name == 'del' or content.name == 'mark' or content.name == 'sub' or content.name == 'sup':
                content_str = content if isinstance(
                    content, str) else str(content)
                if elements and elements[-1]['type'] == 'text':
                    elements[-1]['content'] += content_str
                else:
                    elements.append(
                        {'type': 'text', 'content': content_str})
            elif content.name == 'img':
                elements.append(
                    {'type': 'image', 'content': ("<br/>" + str(content))})
            elif content.name == 'table':
                elements.append(
                    {'type': 'table', 'content': backend_convert_table_to_json(str(content))})
            else:
                parse_content(content.contents)

    soup = BeautifulSoup(text, 'html.parser')
    elements = []

    parse_content(soup.contents)
    return elements

def backend_convert_json_elements_to_html(elements):
    html_content = ''
    for element in elements:
        if element['type'] == 'text':
            formatted_text = element['content'].replace('\n', '<br/>')
            html_content += formatted_text

        elif element['type'] == 'image':
            html_content += element['content']
        elif element['type'] == 'table':
            html_content += backend_convert_table_from_json(element['content'])

    # html_content = html_content.replace('<math>', '<math>')
    html_content = html_content.replace('</math><br/>', '</math> ')
    return html_content

def backend_convert_table_from_json(table_json):
    """Converts a JSON representation of a table into HTML format."""
    table_html = '<table border="1">'
    for row in table_json:
        table_html += '<tr>'
        for cell in row:
            table_html += '<td>'
            # Handle each cell using convert_elements_to_html for recursive processing
            if isinstance(cell, list):
                # If the cell is a list, treat it as a nested table
                table_html += backend_convert_table_from_json(cell)
            elif isinstance(cell, dict) and 'type' in cell:
                # If the cell is a dictionary with a 'type', treat it as an element
                table_html += backend_convert_json_elements_to_html([cell])
            else:
                # Otherwise, treat it as plain text and replace new lines with <br/>
                table_html += str(cell).replace('\n', '<br/>')
            table_html += '</td>'
        table_html += '</tr>'
    table_html += '</table>'
    return table_html

def backend_convert_table_to_json(table_html):
    soup = BeautifulSoup(table_html, 'html.parser')
    table = []
    rows = soup.find_all('tr')
    for row in rows:
        cells = row.find_all(['td', 'th'])
        row_content = []
        for cell in cells:
            cell_html = str(cell).replace('<td>', '').replace(
                '</td>', '').replace('<th>', '').replace('</th>', '')
            row_content.append(cell_html)
        table.append(row_content)
    return table

def convert_docx_to_json(input_file_path, output_file_path):

    def extract_question_data(input_text):
        # Split the text to get the explanation part
        parts = input_text.split("\nExp:")
        explanation = parts[1].strip() if len(parts) > 1 else ""

        # Get the answer part
        parts = parts[0].split("\nAns.")
        answer = parts[1].strip() if len(parts) > 1 else ""

        # Extract options by splitting from the last occurrences of the option labels
        options_text = parts[0]

        # Extract option (d)
        parts = options_text.rsplit("\n(d)", 1)
        option_d = parts[1].strip() if len(parts) > 1 else ""

        # Extract option (c)
        parts = parts[0].rsplit("\n(c)", 1)
        option_c = parts[1].strip() if len(parts) > 1 else ""

        # Extract option (b)
        parts = parts[0].rsplit("\n(b)", 1)
        option_b = parts[1].strip() if len(parts) > 1 else ""

        # Extract option (a) and the question text
        parts = parts[0].rsplit("\n(a)", 1)
        option_a = parts[1].strip() if len(parts) > 1 else ""
        question_text = parts[0].strip()

        # Extract the question number and question
        question_parts = question_text.split(".)", 1)
        question_number = question_parts[0].strip() if len(
            question_parts) > 1 else ""
        question = question_parts[1].strip() if len(question_parts) > 1 else ""

        # Create a JSON object
        data = {
            "question_num": question_number + ".)",
            "question_text": question,
            "options": {
                "a": option_a,
                "b": option_b,
                "c": option_c,
                "d": option_d
            },
            "answer": answer,
            "explanation": explanation
        }

        # return json.dumps(data, ensure_ascii=False, indent=2)
        return data

    def extract_questions(cleaned_html):
        # Split the content into individual question blocks
        question_blocks = re.split(r'(?m)^(\d{1,7}\.\))', cleaned_html)[1:]

        questions = []
        for i in range(0, len(question_blocks), 2):
            question_num = question_blocks[i].strip()
            question_block = question_blocks[i+1].strip()

            question = extract_question_data(question_num + question_block)
            questions.append(question)

        return questions

    def process_questions_with_elements(questions):
        for question in questions:
            # Process question text
            question_elements = backend_extract_html_elements(question['question_text'])
            question['question_elements'] = question_elements

            # Process explanation
            explanation_elements = backend_extract_html_elements(question['explanation'])
            question['explanation_elements'] = explanation_elements

            # Process options
            options_elements = {}
            for key in question['options']:
                option_elements = backend_extract_html_elements(question['options'][key])
                options_elements[key] = option_elements

            question['options_elements'] = options_elements

            # Remove old text keys
            del question['question_text']
            del question['explanation']
            del question['options']

        return questions

    # change only extension of output file path to json  if not already
    if output_file_path.endswith('.docx'):
        output_file_path = output_file_path.replace('.docx', '.json')

    extra_args = [
        '--standalone', 
        '--mathml', 
        '--embed-resources'
    ]

    html_content = pypandoc.convert_file(source_file=input_file_path, to='html5', format='docx', extra_args=extra_args)

    # Clean the HTML content
    cleaned_html = backend_clean_html(html_content)

    # Extract questions from the cleaned HTML content
    questions = extract_questions(cleaned_html)

    # Save the intermediate questions to a JSON file (optional)
    temp_json_file = "temp_json_file.json"

    with open(temp_json_file, 'w', encoding='utf-8') as json_file:
        json.dump(questions, json_file, ensure_ascii=False, indent=4)

    # Load the intermediate questions JSON file
    with open(temp_json_file, 'r', encoding='utf-8') as json_file:
        questions = json.load(json_file)

    # Process the questions to maintain sequence and convert elements
    processed_questions = process_questions_with_elements(questions)

    # Save the final questions to a JSON file
    with open(output_file_path, 'w', encoding='utf-8') as json_file:
        json.dump(processed_questions, json_file, ensure_ascii=False, indent=4)

    if os.path.exists(temp_json_file):
        os.remove(temp_json_file)


    print(
        f"Questions with images, tables, and math formulas have been successfully processed and saved to {output_file_path}")
    return 0

def convert_json_to_docx(input_file_path, output_file_path):


    def create_html_from_json(json_data):
        html_output = '<html><body>'
        for question in json_data:
            html_output += f"{question['question_num']} "
            html_output += backend_convert_json_elements_to_html(
                question['question_elements'])
            html_output += '<p>(a) ' + backend_convert_json_elements_to_html(
                question['options_elements']['a']) + '</p>'
            html_output += '<p>(b) ' + backend_convert_json_elements_to_html(
                question['options_elements']['b']) + '</p>'
            html_output += '<p>(c) ' + backend_convert_json_elements_to_html(
                question['options_elements']['c']) + '</p>'
            html_output += '<p>(d) ' + backend_convert_json_elements_to_html(
                question['options_elements']['d']) + '</p>'
            html_output += f"<p>Ans. {question['answer']}</p>"
            html_output += '<p>Exp: ' + \
                backend_convert_json_elements_to_html(
                    question['explanation_elements']) + '</p>'
        html_output += '</body></html>'
        return html_output


    if not input_file_path.endswith(".json"):
        return 0

    # replace  extension of output file to docx if  not alread
    if output_file_path.endswith(".json"):
        output_file_path = output_file_path.replace(".json", ".docx")

    # Load the JSON data
    with open(input_file_path, 'r', encoding='utf-8', errors='ignore') as json_file:
        questions = json.load(json_file)

    # Create HTML content from JSON data
    html_content = create_html_from_json(questions)

    # Write the HTML content to a temporary HTML file
    output_html_path = "output_html_file.html"
    with open(output_html_path, 'w', encoding='utf-8') as html_file:
        html_file.write(html_content)

    temp_output_docx_file = "temp_output_docx_file.docx"

    extra_args = [
        '--standalone', 
        # '--mathml', 
        # '--embed-resources'
    ]

    pypandoc.convert_file(source_file=output_html_path,outputfile=temp_output_docx_file ,to='docx', format='html', extra_args=extra_args)

    update_font_style_in_docx(temp_output_docx_file, output_file_path)

    if os.path.exists(temp_output_docx_file):
        os.remove(temp_output_docx_file)
    if os.path.exists(output_html_path):
        os.remove(output_html_path)


    print(
        f"DOCX file has been successfully created and saved to {output_file_path}")
    return 0

def remove_newlines_and_extra_spaces(text):

    # remove any html tag from text
    text = re.sub('<.*?>', '', text)
    return ' '.join(text.replace('\n', ' ').split())

def compare_questions(q1, q2):

    # Check if both questions are the same
    if q1 == q2:
        return True   # both questions are the same (ignoring case)

    # Compare all options; assumes both questions have the same option keys (a, b, c, d)
    # for key in ['a', 'b', 'c', 'd']:
    #     q1_option_text = ''.join([elem['content'] for elem in q1['options_elements'][key] if elem['type'] == 'text'])
    #     q2_option_text = ''.join([elem['content'] for elem in q2['options_elements'][key] if elem['type'] == 'text'])
    #     if remove_newlines_and_extra_spaces(q1_option_text) == remove_newlines_and_extra_spaces(q2_option_text) and q1['answer'] == q2['answer'] :
    #         return True
    # if q1['answer'] == q2['answer'] and q1['options_elements'] == q2['options_elements'] and q1['question_elements'] == q2["question_elements"]:
    #     return True

    if q1['answer'] == q2['answer'] and q1['options_elements'] == q2['options_elements']:
        return True
    
    if q1['question_elements'] == q2["question_elements"]:
        return True
    # If all checks passed, the questions are considered the same
    return False

def length_of_explanation(question):

    total_length = 0
    for element in question['explanation_elements']:
        if element['type'] == 'text' or element['type'] == 'math':
            # Add the length of the text content, stripping any extraneous whitespace
            total_length += len(element['content'].strip())
    return total_length

def final_function_to_delete_duplicate_questions(input_folder_path, output_folder_path, folder_with_duplicate):


    def final_function_for_deduplication(folder_path, folder_with_duplicate):

        temp_folder = "temporary"
        ratta_functions.process_folder_for_given_function(
            folder_path, folder_path, convert_docx_to_json, "files")

        json_files = glob.glob(os.path.join(
            folder_path, '**/*.json'), recursive=True)

        duplicates_found = 0
        questions_to_be_deleted_found = 0

        for j, file1 in enumerate(json_files):

            duplicate_question_in_file1 = 0
            question_to_be_deleted_in_file1 = 0

            with open(os.path.join(folder_path, file1), 'r', encoding='utf-8') as f1:
                data1 = json.load(f1)
                duplicate_file_path1 = os.path.join(
                    folder_path, os.path.splitext(file1)[0] + "_D.json")
                print(file1)

                for file2 in json_files[j:]:

                    print(f"        " + str(os.path.basename(file2)))

                    if file1 == file2:
                        for i,  q1 in enumerate(data1):
                            if ratta_functions.is_question_to_be_deleted(q1):
                                if ratta_functions.append_to_json_file(duplicate_file_path1, [q1]):
                                    questions_to_be_deleted_found += 1
                                    question_to_be_deleted_in_file1 += 1
                                    # Logging the process.
                                    # writeLog(f"To be Deleted:")
                                    # writeLog(
                                    # f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                continue
                            for q2 in data1[i+1:]:
                                if compare_questions(q1, q2):

                                    # Logging the process.
                                    # writeLog(f"Found Duplicate in Same File: ")
                                    # writeLog(
                                    #     f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                    # writeLog(
                                    #     f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")

                                    if length_of_explanation(q1) < length_of_explanation(q2):
                                        if ratta_functions.append_to_json_file(duplicate_file_path1, [q1]):
                                            duplicates_found += 1
                                            duplicate_question_in_file1 += 1
                                            # writeLog(  f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                    else:
                                        if ratta_functions.append_to_json_file(duplicate_file_path1, [q2]):
                                            duplicates_found += 1
                                            duplicate_question_in_file1 += 1
                                            # writeLog( f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")
                    else:
                        with open(os.path.join(folder_path, file2), 'r', encoding='utf-8') as f2:
                            data2 = json.load(f2)
                            duplicate_file_path2 = os.path.join(
                                folder_path, os.path.splitext(file2)[0] + "_D.json")

                            for q1 in data1:

                                if ratta_functions.is_question_to_be_deleted(q1):
                                    if ratta_functions.append_to_json_file(duplicate_file_path1, [q1]):
                                        questions_to_be_deleted_found += 1
                                        question_to_be_deleted_in_file1 += 1
                                        # writeLog("To be Deleted: ")
                                        # writeLog(
                                        # f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")

                                    continue

                                for q2 in data2:

                                    if ratta_functions.is_question_to_be_deleted(q2):
                                        if ratta_functions.append_to_json_file(duplicate_file_path2, [q2]):
                                            questions_to_be_deleted_found += 1
                                            # writeLog("To be Deleted: ")
                                            # writeLog(
                                            # f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")
                                        continue

                                    if compare_questions(q1, q2):
                                        # writeLog("Duplicate Found : ")
                                        # writeLog(file1)
                                        # writeLog(
                                        #     f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                        # writeLog(file2)
                                        # writeLog(
                                        # f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")

                                        if length_of_explanation(q1) < length_of_explanation(q2):

                                            if ratta_functions.append_to_json_file(duplicate_file_path1, [q1]):
                                                duplicates_found += 1
                                                duplicate_question_in_file1 += 1
                                                # writeLog(  f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                        else:

                                            if ratta_functions.append_to_json_file(duplicate_file_path2, [q2]):
                                                duplicates_found += 1
                                                # writeLog( f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")
                # print(f"Duplicate: {duplicate_question_in_file1}")
                print(
                    f"Questions to be deleted: {question_to_be_deleted_in_file1}")
        for file_path in json_files:
            os.remove(file_path)

        ratta_functions.process_folder_for_given_function(
            folder_path, folder_with_duplicate, convert_json_to_docx, "files")

        # json_files = glob.glob(os.path.join(
        #     folder_path, '**/*.json'), recursive=True)
        # for file_path in json_files:
        #     os.remove(file_path)
        if os.path.exists(temp_folder):
            shutil.rmtree(temp_folder)

        return duplicates_found, questions_to_be_deleted_found

    def delete_duplicate_question(input_file, output_file):

        try:
            if input_file.lower().endswith(".json"):
                return 0

            duplicate_json_file = os.path.basename(
                input_file).replace('.docx', '_D.json')
            duplicate_json_file = os.path.join(
                os.path.dirname(input_file), duplicate_json_file)

            if not os.path.exists(duplicate_json_file):
                shutil.copy(input_file, output_file)
                return 0

            original_input_json_file = "original_input_json_file.json"
            convert_docx_to_json(input_file, original_input_json_file)

            with open(original_input_json_file, 'r', encoding="utf-8", errors='ignore') as file:
                original_data = json.load(file)

            with open(duplicate_json_file, 'r', encoding="utf-8", errors='ignore') as file:
                duplicate_data = json.load(file)

            os.remove(original_input_json_file)
            # delete duplicate json object from original_data
            filtered_data = [
                obj for obj in original_data if obj not in duplicate_data]

            # Get the base name without extension
            temporary_json_file_with_filtered_data = "temporary_json_file_with_filtered_data.json"
            with open(temporary_json_file_with_filtered_data, 'w', encoding='utf-8') as json_file:
                json.dump(filtered_data, json_file,
                          ensure_ascii=False, indent=4)

            convert_json_to_docx(
                temporary_json_file_with_filtered_data, output_file)

            os.remove(temporary_json_file_with_filtered_data)

        except Exception as e:
            ratta_functions.writeLog(f"An error occurred: {e}")
        return 0

    if os.path.exists(folder_with_duplicate):
        shutil.rmtree(folder_with_duplicate)

    if os.path.exists(output_folder_path):
        shutil.rmtree(output_folder_path)

    duplicates_found, questions_to_be_deleted_found = final_function_for_deduplication(
        input_folder_path, folder_with_duplicate)

    ratta_functions.process_folder_for_given_function(
        input_folder_path, output_folder_path, delete_duplicate_question, "files")

    json_files = glob.glob(os.path.join(
        input_folder_path, '**/*.json'), recursive=True)
    for file_path in json_files:
        os.remove(file_path)

    return duplicates_found, questions_to_be_deleted_found

def get_initial_idea_of_files_from_a_folder_recursively(folder_path):

    def aggregate_stats(old_stats, new_stats):
        total_questions = old_stats["Total Questions"] + \
            new_stats["Total Questions"]
        total_no_explanation = old_stats["No Explanation"] + \
            new_stats["No Explanation"]
        if total_questions > 0:
            percentage_no_explanation = round(
                (total_no_explanation / total_questions) * 100, 1)
        else:
            percentage_no_explanation = 0

        if total_questions - total_no_explanation > 0:
            explanation_length = round((old_stats["Explanation length"] * (old_stats["Total Questions"] - old_stats["No Explanation"]) +
                                        new_stats["Explanation length"] * (new_stats["Total Questions"] - new_stats["No Explanation"])) /
                                       (total_questions - total_no_explanation), 1)
        else:
            explanation_length = 0

        return {
            "Total Questions": total_questions,
            "No Explanation": total_no_explanation,
            "Question length": round((old_stats["Question length"] * old_stats["Total Questions"] +
                                      new_stats["Total Questions"] * new_stats["Question length"]) / total_questions, 1) if total_questions > 0 else 0,
            "Percent NO Explanation": percentage_no_explanation,
            "Option length": round((old_stats["Option length"] * old_stats["Total Questions"] +
                                    new_stats["Total Questions"] * new_stats["Option length"]) / total_questions, 1) if total_questions > 0 else 0,
            "Explanation length": explanation_length
        }

    def extract_stats_from_file(input_file):
        with open(input_file, 'r', encoding='utf-8') as f:
            data = json.load(f)

        length_of_question = 0
        length_of_option = 0
        length_of_explanation = 0
        number_of_question = len(data)
        number_of_question_not_having_explanation = 0

        for item in data:
            # Calculate the length of questions
            for element in item['question_elements']:
                if element['type'] == 'text':
                    length_of_question += len(element['content'].split())

            # Calculate the length of options
            for key in ['a', 'b', 'c', 'd']:
                for element in item['options_elements'][key]:
                    if element['type'] == 'text':
                        length_of_option += len(element['content'].split())

            # Calculate the length of explanations
            explanation_length = 0
            for element in item['explanation_elements']:
                if element['type'] == 'text':
                    explanation_length += len(element['content'].split())

            if explanation_length == 0:
                number_of_question_not_having_explanation += 1
            else:
                length_of_explanation += explanation_length

        if number_of_question > 0:
            average_of_length_of_question = round(
                length_of_question / number_of_question, 1)
            average_of_length_of_options = round(
                length_of_option / (number_of_question * 4), 1)
            percentage_of_question_not_having_explanation = round(
                (number_of_question_not_having_explanation / number_of_question) * 100, 1)
            if (number_of_question - number_of_question_not_having_explanation) > 0:
                average_of_length_of_explanation = round(
                    length_of_explanation / (number_of_question - number_of_question_not_having_explanation), 1)
            else:
                # if all questions have an explanation, then the length of explanations is zero.
                average_of_length_of_explanation = 0
        else:
            average_of_length_of_question = 0
            average_of_length_of_options = 0
            percentage_of_question_not_having_explanation = 0
            average_of_length_of_explanation = 0

        return {
            "Total Questions": number_of_question,
            "No Explanation": number_of_question_not_having_explanation,
            "Percent NO Explanation": percentage_of_question_not_having_explanation,
            "Question length": average_of_length_of_question,
            "Option length": average_of_length_of_options,
            "Explanation length": average_of_length_of_explanation
        }

    def extract_stats_from_folder(folder_path):
        folder_data = {"Folder": os.path.basename(folder_path)}
        if os.path.isdir(folder_path):
            folder_stats = {
                "Total Questions": 0,
                "Percent NO Explanation": 0,
                "No Explanation": 0,
                "Question length": 0,
                "Option length": 0,
                "Explanation length": 0
            }
            children = []
            for item in sorted(os.listdir(folder_path)):
                item_path = os.path.join(folder_path, item)
                if os.path.isdir(item_path):
                    child_data, child_stats = extract_stats_from_folder(
                        item_path)
                    children.append(child_data)
                    folder_stats = aggregate_stats(folder_stats, child_stats)
                elif item.endswith('.json'):
                    stats = extract_stats_from_file(item_path)
                    children.append({"File": item, **stats})
                    folder_stats = aggregate_stats(folder_stats, stats)

            folder_data.update(folder_stats)
            folder_data["Children"] = children
        return folder_data, folder_stats

    temp_folder = "temporary"
    another_temp_folder = os.path.basename(folder_path)

    ratta_functions.process_folder_for_given_function(
        folder_path, another_temp_folder, convert_docx_to_json, "files")

    json_output = os.path.basename(folder_path) + ".json"
    text_output = os.path.basename(folder_path) + ".text"
    folder_structure, _ = extract_stats_from_folder(another_temp_folder)
    ratta_functions.write_json(folder_structure, json_output)
    ratta_functions.make_json_readable(json_output, text_output)

    # delete temp folder
    if os.path.exists(temp_folder):
        shutil.rmtree(temp_folder)

    if os.path.exists(another_temp_folder):
        shutil.rmtree(another_temp_folder)

def segregate_question(input_folder_path, output_folder_path, guide_line_json, search_in_explanation=False):

    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)

    temp_folder = "temporary"
    another_temp = "another_temp"
    ratta_functions.process_folder_for_given_function(
        input_folder_path, input_folder_path, convert_docx_to_json, "files")

    json_files = glob.glob(os.path.join(
        input_folder_path, '**/*.json'), recursive=True)

    for filename in json_files:
        basename = os.path.basename(filename)
        with open(os.path.join(input_folder_path, filename), 'r') as f:
            data = json.load(f)
            for question in data:
                print(question)
                # question_content =    ''.join([elem['content'] for elem in question['question_elements'] if elem['type'] == 'text' or elem['type'] =='table'])
                question_content = str(question['question_elements'])

                # explanation = ''.join([ elem['content'] for elem in question['explanation_elements'] if elem['type'] == 'text' or elem['type'] == 'table'])
                explanation = str(question['explanation_elements'])
                # answer_content =''.join([elem['content'] for elem in question['options_elements'][question['answer']] if elem['type'] == 'text' or elem['type'] == 'table'])
                answer_content = str(
                    question['options_elements'][question['answer']])
                matched = False
                for key, values in guide_line_json.items():
                    for value in values:
                        if (value in question_content) or (value in answer_content):
                            ratta_functions.append_to_json_file(os.path.join(
                                output_folder_path, key + '.json'), [question])
                            matched = True
                            break
                    if matched:
                        break
                if not matched:
                    if search_in_explanation:
                        for key, values in guide_line_json.items():
                            for value in values:
                                if (value in explanation):
                                    ratta_functions.append_to_json_file(os.path.join(
                                        output_folder_path, key + '.json'), [question])
                                    matched = True
                                    break
                            if matched:
                                break
                if not matched:
                    ratta_functions.append_to_json_file(
                        os.path.join(output_folder_path, basename), [question])
    temp = "jasdkjfakljalksjdg"
    ratta_functions.process_folder_for_given_function(
        output_folder_path, output_folder_path, convert_json_to_docx, "files")

    txt_files = glob.glob(os.path.join(
        output_folder_path, '**/*.json'), recursive=True)
    for file_path in txt_files:
        os.remove(file_path)

    for file_path in json_files:
        os.remove(file_path)
    if os.path.exists(temp_folder):
        shutil.rmtree(temp_folder)

    if os.path.exists(temp):
        shutil.rmtree(temp)

def find_wrong_format_question(input_file_path, output_file_path):

    temp_file = "test.json"
    wrong_format_json_file = "wrong_format_json_file.json"

    convert_docx_to_json(input_file_path,  temp_file)

    if "_WRONG" not in output_file_path:
        output_file_path = os.path.splitext(output_file_path)[
            0] + "_WRONG" + os.path.splitext(output_file_path)[1]

    with open(temp_file, 'r') as f:
        data = json.load(f)
        for question in data:
            question_content = ''.join(
                [elem['content'] for elem in question['question_elements'] if elem['type'] == 'text' or elem['type'] == 'image'])
            # table_exist = "".join( "Table exist " if elem['type'] == 'table' else "" for elem in question['question_elements']  )
            answer = question['answer']
            # answer_content = ''.join([elem['content'] for elem in question['options_elements'][answer] ] )
            option_a = ''.join([elem['content'] for elem in question['options_elements']
                               ['a'] if elem['type'] == 'text' or elem['type'] == 'image'])
            option_b = ''.join([elem['content'] for elem in question['options_elements']
                               ['b'] if elem['type'] == 'text' or elem['type'] == 'image'])
            option_c = ''.join([elem['content'] for elem in question['options_elements']
                               ['c'] if elem['type'] == 'text' or elem['type'] == 'image'])
            option_d = ''.join([elem['content'] for elem in question['options_elements']
                               ['d'] if elem['type'] == 'text' or elem['type'] == 'image'])

            if answer not in ['a', 'b', 'c', 'd']:
                ratta_functions.append_to_json_file(
                    wrong_format_json_file, [question])
            elif option_a == "" or option_b == "" or option_c == "" or option_d == "":
                ratta_functions.append_to_json_file(
                    wrong_format_json_file, [question])
            elif question_content == "":
                ratta_functions.append_to_json_file(
                    wrong_format_json_file, [question])

    if os.path.exists(wrong_format_json_file):
        convert_json_to_docx(wrong_format_json_file, output_file_path)

    if os.path.exists(temp_file):
        os.remove(temp_file)
    if os.path.exists(wrong_format_json_file):
        os.remove(wrong_format_json_file)
    return 0

def make_question_readable(obj):
    def format_element(element):
        if 'type' in element and element['type'] == 'text':
            return element['content']
        elif 'type' in element and element['type'] == 'table':
            rows = []
            for row in element['content']:
                rows.append(' | '.join(row))
            return '\n'.join(rows)
        if 'type' in element and element['type'] == 'image':
            # content = element['content']
            # base64_string = content.split('base64,')[1].split('\\"')[0]
            # img_data = base64.b64decode(base64_string)
            # img = Image.open(BytesIO(img_data))
            # img.save("image.png")
            # return content.split('<img')[0]
            return "some Image here"

    question_num = obj['question_num']
    answer = obj['answer']
    question = '\\n'.join([format_element(e) for e in obj['question_elements']])
    explanation = '\\n'.join([format_element(e) for e in obj['explanation_elements']])
    options = '\n'.join([f'{key}: {format_element(option[0])}' for key, option in obj['options_elements'].items()])

    formatted_string = f"""
    {question_num} {question}
    \n Options:
    {options}
    \n Answer: {answer}
    \n Explanation:
    {explanation}
    """
    return formatted_string

def integrate_new_folder_into_existing_folder(existing_folder_path, new_folder_path, output_folder_path):
    if os.path.exists(output_folder_path):
        shutil.rmtree(output_folder_path)

    os.makedirs(output_folder_path)
    

    existing_folder_json = "existing_folder_json"
    new_folder_json = "new_folder_json"
    output_folder_json = "output_folder_json"
    duplicate_questions_json = "duplicate_questions.json"
    duplicate_questions_file = "duplicate_questions.docx"


    ratta_functions.process_folder_for_given_function(existing_folder_path,existing_folder_json,convert_docx_to_json, "files")
    ratta_functions.process_folder_for_given_function(new_folder_path, new_folder_json, convert_docx_to_json, "files")

    new_json_files = glob.glob(os.path.join(new_folder_json, '**/*.json'), recursive=True)
    existing_json_files = glob.glob(os.path.join(existing_folder_json, '**/*.json'), recursive=True)

    combined_json = "combined_json"

    combined_existing_data = []

    for file in existing_json_files:
        with open(file, 'r', encoding='utf-8') as f1:
            data1 = json.load(f1)
        combined_existing_data.extend(data1)


    for file in new_json_files:

        print(file)
        new_file = file.replace(new_folder_json, output_folder_json)
        new_questions_to_add = []

        duplicate_questions = []

        with open(file, 'r', encoding='utf-8') as f1:
            data1 = json.load(f1)
        for new_question in data1:
            if not any(compare_questions(new_question, old_question) for old_question in combined_existing_data):
                new_questions_to_add.append(new_question)
            else:
                duplicate_questions.append(new_question)
                # duplicate_questions.append(old_question)


            
        if new_questions_to_add:
            ratta_functions.append_to_json_file(new_file, new_questions_to_add)
        if duplicate_questions:
            ratta_functions.append_to_json_file(duplicate_questions_json,duplicate_questions)



    ratta_functions.process_folder_for_given_function(output_folder_json, output_folder_path, convert_json_to_docx, "files")
    convert_json_to_docx(duplicate_questions_json, duplicate_questions_file)

    if os.path.exists(existing_folder_json):
        shutil.rmtree(existing_folder_json)
    if os.path.exists(new_folder_json):
        shutil.rmtree(new_folder_json)
    if os.path.exists(output_folder_json):
        shutil.rmtree(output_folder_json)

def get_cummulative_explanation(input_file_path , output_file_path):

    temp_json_file = "temp_extra_json_file.json"
    temp_hmtl_output = "temp_extra_hmtl_output.html"
    temp_output_docx_file = "temp_extra_output.docx"
    
    explanation_array = []

    # explanation_string = ""
    convert_docx_to_json(input_file_path, temp_json_file)

    with open(temp_json_file, 'r') as f:
        data = json.load(f)

    for question in data:
        explanation_array.extend(question['explanation_elements'])
        explanation_array.extend([{'type':'text', 'content':'\n'}])
 
    html_content  =  backend_convert_json_elements_to_html(explanation_array)

    # save explanation_string to output_file_path
    with open(temp_hmtl_output, 'w') as f:
        f.write(html_content)


    extra_args = [
        '--standalone', 
        # '--mathml', 
        # '--embed-resources'
    ]

    pypandoc.convert_file(source_file=temp_hmtl_output,outputfile=temp_output_docx_file ,to='docx', format='html', extra_args=extra_args)

    update_font_style_in_docx(temp_output_docx_file, output_file_path)

    if os.path.exists(temp_json_file):
        os.remove(temp_json_file)
    if os.path.exists(temp_hmtl_output):
        os.remove(temp_hmtl_output)
    if os.path.exists(temp_output_docx_file):
        os.remove(temp_output_docx_file)

    return 0

def convert_any_docx_file_to_json_file(input_file_path, output_file_path):

    # temp_extra_explanation_html_file = "temp_extra_explanation.html"
    temp_extra_explanation_json_file = "temp_extra_explanation.json"


    extra_args = [
        '--standalone', 
        '--mathml', 
        '--embed-resources'
    ]

    html_extra_content = pypandoc.convert_file(source_file=input_file_path, to='html5', format='docx', extra_args=extra_args)



    # Use regular expressions to extract content within the <body> tags
    match = re.search(r'<body[^>]*>(.*?)</body>', html_extra_content, re.DOTALL | re.IGNORECASE)

    # Check if a match was found and extract the body content
    if match:
        body_content = match.group(1)
    else:
        body_content = ''  # If no <body> tag is found, set body_content to an empty string

    # Clean the HTML content

    cleaned_html = backend_clean_html(body_content)

    extra_elements = backend_extract_html_elements(cleaned_html)

    with open(output_file_path, 'w') as file:
        json.dump(extra_elements, file, ensure_ascii=False, indent=4)


    if os.path.exists(temp_extra_explanation_json_file):
        os.remove(temp_extra_explanation_json_file)

    return 0

def append_extra_explanation_to_every_question_file(input_file_path, output_file_path):


    extra_elements_file = input_file_path.replace("input" , "extra")
    extra_elements_json_file = extra_elements_file.replace(".docx", "_extra_temp.json")

    convert_any_docx_file_to_json_file(extra_elements_file, extra_elements_json_file)

    with open(extra_elements_json_file, 'r') as f:
        extra_elements_to_append = json.load(f)

    questions_json_file = input_file_path.replace(".docx" , "_questions_temp.json")
    convert_docx_to_json(input_file_path, questions_json_file)  

    with open(questions_json_file, 'r') as f:
        data1 = json.load(f)

    for question in data1:
        question['explanation_elements'].extend(extra_elements_to_append)


    with open(questions_json_file, 'w') as f:
        json.dump(data1, f, ensure_ascii=False, indent=4)

    convert_json_to_docx(questions_json_file, output_file_path)

    if os.path.exists(extra_elements_json_file):
        os.remove(extra_elements_json_file)
    if os.path.exists(questions_json_file):
        os.remove(questions_json_file)

    return 0

def convert_docx_to_markdown(input_file_path, output_file_path):

    if not output_file_path.endswith('.md'):
        output_file_path = output_file_path.replace("docx" , ".md")
    try:
        # Call pandoc via subprocess to convert the DOCX to Markdown
        # subprocess.run(
        #     ['pandoc', input_file_path, '-t', 'markdown', '-o', output_file_path],
        #     check=True
        # )
        pypandoc.convert_file(source_file=input_file_path, to='markdown', format='docx', outputfile=output_file_path)

        print(f"Conversion successful! Markdown saved to: {output_file_path}")
    except subprocess.CalledProcessError as e:
        print(f"An error occurred during conversion: {e}")
    return 0

def convert_markdown_to_docx(input_file_path, output_file_path):
    # Ensure the output file has the correct .docx extension
    if not output_file_path.endswith('.docx'):
        output_file_path = output_file_path.replace(".md", ".docx")
    
    try:
        # Call pandoc via subprocess to convert the Markdown to DOCX
        # subprocess.run(
        #     ['pandoc', input_file_path, '-o', output_file_path],
        #     check=True
        # )
        pypandoc.convert_file(source_file=input_file_path, to='docx', format='markdown', outputfile=output_file_path)
        
        print(f"Conversion successful! DOCX saved to: {output_file_path}")
    except subprocess.CalledProcessError as e:
        print(f"An error occurred during conversion: {e}")
    
    return 0

def convert_text_string_to_json(string_to_convert):
 
    
    def extract_question_data_from_string(input_text):
        parts = input_text.split("\nExp:")
        explanation = parts[1].strip() if len(parts) > 1 else ""

        parts = parts[0].split("\nAns.")
        answer = parts[1].strip() if len(parts) > 1 else ""

        options_text = parts[0]

        parts = options_text.rsplit("\n(d)", 1)
        option_d = parts[1].strip() if len(parts) > 1 else ""

        parts = parts[0].rsplit("\n(c)", 1)
        option_c = parts[1].strip() if len(parts) > 1 else ""

        parts = parts[0].rsplit("\n(b)", 1)
        option_b = parts[1].strip() if len(parts) > 1 else ""

        parts = parts[0].rsplit("\n(a)", 1)
        option_a = parts[1].strip() if len(parts) > 1 else ""
        question_text = parts[0].strip()

        question_parts = question_text.split(".)", 1)
        question_number = question_parts[0].strip() + ".)" if len(question_parts) > 1 else ""
        question = question_parts[1].strip() if len(question_parts) > 1 else ""

        return {
            "question_num": question_number,
            "question_text": question,
            "options": {
                "a": option_a,
                "b": option_b,
                "c": option_c,
                "d": option_d
            },
            "answer": answer,
            "explanation": explanation
        }

    # Clean the input string to ensure proper formatting
    cleaned_string = remove_newlines_and_extra_spaces(string_to_convert)
    
    # Extract question data
    question_data = extract_question_data_from_string(cleaned_string)
    
    # Process question elements
    question_elements = backend_extract_html_elements(question_data['question_text'])
    question_data['question_elements'] = question_elements

    # Process explanation
    explanation_elements = backend_extract_html_elements(question_data['explanation'])
    question_data['explanation_elements'] = explanation_elements

    # Process options
    options_elements = {}
    for key in question_data['options']:
        option_elements = backend_extract_html_elements(question_data['options'][key])
        options_elements[key] = option_elements

    question_data['options_elements'] = options_elements

    # Remove old text keys
    del question_data['question_text']
    del question_data['explanation']
    del question_data['options']

    return question_data

def segregate_question_using_number(input_file_path, output_folder_path, question_map):

    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)

    original_input_json_file = input_file_path.replace(".docx", ".json")

    convert_docx_to_json(input_file_path, original_input_json_file)

    with open(original_input_json_file, 'r') as f:
        data1 = json.load(f)

    
        for question in data1:

                print(question)
                question_number_text = question['question_num'].split(".")[0]
                
                question_number = int(question_number_text)


               
                matched = False
                for key, values in question_map.items():
                    for value in values:
                        if (value == question_number) :
                            ratta_functions.append_to_json_file(os.path.join(
                                output_folder_path, key + '.json'), [question])
                            matched = True
                            break
                    if matched:
                        break
                
                if not matched:
                    ratta_functions.append_to_json_file(
                        os.path.join(output_folder_path, 'others.json'), [question])

    ratta_functions.process_folder_for_given_function(
        output_folder_path, output_folder_path, convert_json_to_docx, "files")

    txt_files = glob.glob(os.path.join(
        output_folder_path, '**/*.json'), recursive=True)
    for file_path in txt_files:
        os.remove(file_path)

    if os.path.exists(original_input_json_file):
        os.remove(original_input_json_file)

def convert_excel_to_json(input_file, output_file):
    input_file = os.path.expanduser(input_file)
    output_file = os.path.expanduser(output_file)

    # Initialize a dictionary to hold all sheets' data
    json_data = {}
    
    try:
        # Read the Excel file without headers (header=None)
        df = pd.read_excel(input_file, sheet_name=None, header=None)
        
        # Iterate through each sheet
        for sheet_name, data in df.items():
            # Assign column names like column_1, column_2, ...
            data.columns = [f'column_{i+1}' for i in range(data.shape[1])]
            # Convert the DataFrame to a list of dictionaries
            json_data[sheet_name] = data.to_dict(orient='records')
        
        # Write the initial JSON data to a temporary file
        temp_json_file = output_file.replace('.json', '_temp.json')
        with open(temp_json_file, 'w', encoding='utf-8') as outfile:
            json.dump(json_data, outfile, ensure_ascii=False, indent=4)

    except Exception as e:
        print(f"Error reading Excel file: {e}")
        return

    # Now read the temporary JSON file and process it
    output = {}

    try:
        with open(temp_json_file, 'r', encoding='utf-8') as infile:
            data = json.load(infile)

        # Process the JSON data
        for sheet_name, entries in data.items():
            for entry in entries:
                # Extract the key and values
                key = entry.get('column_1')
                values = entry.get('column_2')

                # Handle NaN and convert it to an empty list
                if pd.isna(values):
                    value_list = []
                else:
                    if isinstance(values, str):
                        # Split the string and convert to a list of integers
                        value_list = [int(x.strip()) for x in values.split(',') if x.strip().isdigit()]
                    else:
                        value_list = [int(values)]

                # Assign to the output dictionary
                output[key] = value_list

        # Write the processed output to the final JSON file
        with open(output_file, 'w', encoding='utf-8') as outfile:
            json.dump(output, outfile, ensure_ascii=False, indent=4)

    except Exception as e:
        print(f"Error processing JSON file: {e}")
        return

    # Optionally, remove the temporary file if needed
    os.remove(temp_json_file)

    return output

def convert_answer_docx_to_json(input_file_path, output_file_path):
    
    if output_file_path.endswith('.docx'):
        output_file_path = output_file_path.replace('.docx', '.json')

    extra_args = [
        '--standalone', 
        '--mathml', 
        '--embed-resources'
    ]

    html_content = pypandoc.convert_file(source_file=input_file_path, to='html5', format='docx', extra_args=extra_args)

    cleaned_html = backend_clean_html(html_content)

    question_blocks = re.split(r'(?m)^(\d{1,7}\.\))', cleaned_html)[1:]

    questions = []
    for i in range(0, len(question_blocks), 2):
        question_num = question_blocks[i].strip()
        question_block = question_blocks[i+1].strip()
        data = {
            'question_num': question_num,
            'explanation_elements': backend_extract_html_elements(question_block)
        }
        
        questions.append(data)

    with open(output_file_path, 'w', encoding='utf-8') as json_file:
        json.dump(questions, json_file, ensure_ascii=False, indent=4)

def merge_question_answer_docx(input_file_path,answer_docx_file, output_file_path):

    # if output file path folder not exist then create
    if not os.path.exists(os.path.dirname(output_file_path)):
        os.makedirs(os.path.dirname(output_file_path))

    # answer_docx_file = input_file_path.replace("input" , "extra")
    temp_answer_json = answer_docx_file.replace(".docx", ".json")
    temp_question_json = input_file_path.replace(".docx", ".json")
    temp_output_json = output_file_path.replace(".docx", ".json")

    convert_answer_docx_to_json(answer_docx_file, temp_answer_json)

    convert_docx_to_json(input_file_path, temp_question_json)

    with open(temp_question_json, 'r') as f:
        question_data = json.load(f)

    with open(temp_answer_json, 'r') as f:
        answer_data = json.load(f)


    # take answer from answer_data and put in question_data

    for question in question_data:
        for answer in answer_data:
            if question['question_num'] == answer['question_num']:
                question['explanation_elements'] = answer['explanation_elements']                
                break

    with open(temp_output_json, 'w', encoding='utf-8') as f:
        json.dump(question_data, f, ensure_ascii=False, indent=4)

    convert_json_to_docx(temp_output_json, output_file_path)

    if os.path.exists(temp_answer_json):
        os.remove(temp_answer_json)
    if os.path.exists(temp_question_json):
        os.remove(temp_question_json)
    if os.path.exists(temp_output_json):
        os.remove(temp_output_json)


# import os
# import logging
# from pathlib import Path

# def process_folder_for_given_function(input_folder, output_folder, process_function, file_or_folder="files", allowed_extensions=None):
#     """
#     Recursively processes files in the input folder, maintains folder structure, and applies the
#     specified process function to each file, saving the processed files in the output folder.
    
#     Parameters:
#     - input_folder: Path to the folder containing files to process.
#     - output_folder: Path where processed files will be saved.
#     - process_function: Function to apply to each file (input path, output path).
#     - file_or_folder: "files" to process files, "folders" to process directories.
#     - allowed_extensions: Optional set of file extensions to include in processing.
    
#     Returns:
#     Total count of processed files or folders.
#     """
    
#     logging.basicConfig(level=logging.INFO)
#     total_processed = 0
    
#     input_folder = Path(input_folder)
#     output_folder = Path(output_folder)

#     if not input_folder.is_dir():
#         raise ValueError("input_folder must be a valid directory")

#     output_folder.mkdir(parents=True, exist_ok=True)

#     for root, directories, files in os.walk(input_folder):
#         relative_path = Path(root).relative_to(input_folder)
#         output_subfolder = output_folder / relative_path

#         output_subfolder.mkdir(parents=True, exist_ok=True)

#         if file_or_folder == "files":
#             for filename in files:
#                 input_file_path = Path(root) / filename
#                 output_file_path = output_subfolder / filename
                
#                 # Convert filename to string for checking extensions
#                 filename_str = str(filename)
                
#                 # Check for allowed extensions
#                 if allowed_extensions is None or any(filename_str.endswith(ext) for ext in allowed_extensions):
#                     try:
#                         total_processed += process_function(input_file_path, output_file_path)
#                         logging.info(f"Processed file: {filename_str}")
#                     except Exception as e:
#                         logging.error(f"Error processing {input_file_path}: {e}")
#                 else:
#                     logging.info(f"Skipped file: {filename_str} (not in allowed extensions)")

#         else:
#             try:
#                 total_processed += process_function(root, output_subfolder)
#             except Exception as e:
#                 logging.error(f"Error processing folder {root}: {e}")

#     return total_processed


