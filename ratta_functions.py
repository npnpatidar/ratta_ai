import json
import os
import yaml
import re
import csv
import random
# import g4f
import docx2txt
import pypandoc
import docx
from itertools import combinations
import shutil
import glob
from graphviz import Digraph
import xml.etree.ElementTree as ET
import logging
from intelligence import Brain
import time

from docx import Document

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Inches

system_instruction = """I will provide a question, answer and explanation which has some mistakes
due to bad OCR . I want you to correct it and provide proper json object for it in below format.
remember these:
1. the questions should always start with  "1.)" 
2. If explanation is blank then provide good but short and precise explanation but if it is given don't change it.
3. If answer is blank then provide the correct answer.
4. if some option is missing then provide apporpriate option.

  {
    "question": "1.) A basic unit of measurement of information storage in a computer.",
    "option_a": "gram",
    "option_b": "meter",
    "option_c": "Byte",
    "option_d": "bit per second",
    "answer": "c",
    "answer_content": "Byte",
    "explanation": "Byte = Smallest Unit of measurement of information srorage in computer."
  }
"""
# microBrain = Brain(system_instruction=system_instruction)


def compare_two_questions_with_ai(q1, q2):
    is_duplicate = False
    system_instruction = """ I have done OCR on two papers of the same exam but different paper set.
    The question belong to same exam so the language of the question will be same but options may be interachanged.
    All options of question may be similar but their order can be different.
    Since I have done OCR using bad scanner there may be lots of mistakes in OCR.
    you can ignore slight variations in the structure of the questions.
    you can ignore minor variations and typogrphical mistakes.
    there can be unintentional spaces or unrelated character in the questions due to bad OCR
    if the options are mostly similar then consider those questions as same.
    You can ignore explanation given in the question for the answer.
    you need to compare them and decide whether these questions are  the same or not.
    if the questions are SAME then your response should be TRUE else FALSE.
    Your response MUST be in JSON in below format.
    The reason why the questions are same or different is explained in "reason" field.
    If you return anything except the JSON format an innocent kitten will be killed horribly.
    you need to save kittens by giving response in proper JSON format as given below.

    {
     "isDuplicate" : "TRUE",
     "reason": "why the questions are same or different"
    }

    """

    question_1 = json.dumps(q1)
    question_2 = json.dumps(q2)
    questions = question_1 + "\n\n" + question_2
    answer = microBrain.ask_question(system_instruction=system_instruction, question=questions,
                                     host="openrouter", model="mistralai/mistral-7b-instruct:free", json_format=True)
    answer = answer[answer.index("{"): answer.index("}")+1]
    writeLog(answer)
    json_answer = json.loads(answer)
    time.sleep(2)
    writeLog(json_answer)
    print(json_answer)
    if json_answer["isDuplicate"] == "TRUE":
        is_duplicate = True
    return is_duplicate


def setup_logger(log_file='process.log'):
    logging.basicConfig(filename=log_file, level=logging.INFO,
                        format='%(asctime)s - %(levelname)s - %(message)s')


# def writeLog(message):
#     logging.info(message)


def writeLog(message):
    with open("process.log", "a") as f:
        f.write("\n")
        f.write(str(message))


def backend_folder_to_json(folder_path, include_files=False):
    """
    This function takes the path of a folder and returns a dictionary
    representing its folder tree structure, optionally including files.

    Args:
        folder_path (str): The path to the folder.
        include_files (bool, optional): Whether to include files in the output. Defaults to False.

    Returns:
        dict: A dictionary representing the folder tree structure.
    """
    # Check if the path is a directory
    if not os.path.isdir(folder_path):
        raise ValueError(f"{folder_path} is not a directory")

    # Initialize the result dictionary
    result = {"Folder": os.path.basename(folder_path)}

    # Handle subfolders and (optionally) files
    children = []
    # Sort entries alphabetically
    entries = sorted(os.listdir(folder_path), key=lambda x: x.lower())
    for entry in entries:
        full_path = os.path.join(folder_path, entry)
        if os.path.isdir(full_path):
            children.append(backend_folder_to_json(full_path, include_files))
        elif include_files:
            children.append({"File": entry})

    # Add children to the result dictionary
    if children:
        result["Children"] = children

    return result


def save_folder_structure_into_yaml(folder_path, output_folder, include_files=False):
    """
    This function saves a dictionary to a JSON file.

    Args:
        filename (str): The name of the file to save to.
        data (dict): The data to save.

    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    data = backend_folder_to_json(folder_path, include_files)

    json_file_path = os.path.join(output_folder, "folder_structure.json")
    yaml_file_path = os.path.join(output_folder, "folder_structure.yaml")
    with open(json_file_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

    with open(yaml_file_path, 'w') as f:
        yaml.dump(data, f, allow_unicode=True, indent=8,
                  default_flow_style=False, sort_keys=False)


def backend_create_folders_and_files(folder_structure, parent_path):
    """
    This function creates folders and files based on the provided folder structure data.

    Args:
        folder_structure (dict): The dictionary representing the folder structure.
        parent_path (str): The path where the folders and files will be created.
    """
    folder_name = folder_structure["Folder"]
    current_path = os.path.join(parent_path, folder_name)

    # Create folder if it doesn't exist
    if not os.path.exists(current_path):
        os.makedirs(current_path)
        print("created folder: ", current_path)

    # Check if there are children
    if "Children" in folder_structure:
        children = folder_structure["Children"]
        for child in children:
            if "File" in child:
                file_name = child["File"]
                file_path = os.path.join(current_path, file_name)
                # Create empty file
                # open(file_path, 'a').close()
            else:
                # Recursively create child folders and files
                backend_create_folders_and_files(child, current_path)


def create_folder_structure_from_yaml(yaml_file, output_folder):
    """
    This function creates folder structure based on the YAML file.

    Args:
        yaml_file (str): The path to the YAML file containing folder structure data.
        output_folder (str): The path where the folders and files will be created.
    """
    with open(yaml_file, 'r') as file:
        folder_structure = yaml.safe_load(file)

    backend_create_folders_and_files(folder_structure, output_folder)


def chatgpt(prompt):

    client = OpenAI(
        # This is the default and can be omitted
        api_key=os.environ.get("OPENAI_API_KEY"),
    )

    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt,
            }
        ],
        model="gpt-3.5-turbo",
    )


def call_g4f_api(prompt):
    max_attempts = 1
    reformatted_question = ""
    # Define your conversation with the model
    conversation = [
        {
            "role":
            "system",
            "content":
            prompt
        },
    ]
    for _ in range(max_attempts):
        try:
            response = g4f.ChatCompletion.create(
                model="gpt-4",
                messages=conversation,
                max_tokens=2000,
                stream=False,
            )

            for message in response:
                reformatted_question += message

            # Split the response into words and check if it has more than 5 words
            words = reformatted_question.split()
            if len(words) > 0:
                return reformatted_question

        except Exception as e:
            # Log the error (you can use a logging library for this)
            print(f"Error while reformatting question: {str(e)}")
            # print(f"error in summarising  ")

    # If after 10 attempts there's no valid response, return an error message or handle as needed
    return None


def g4f_api(input_file, output_file):

    # Call the function and get the array of questions
    questions_array = extract_questions(input_file)

    prompt_for_each_question = "Just translate this in english and keept the same format"
    # Print or use the array as needed
    for i, question in enumerate(questions_array, start=1):
        # print(f"{i}. {question}")

        reformatted_question = call_g4f_api(
            prompt_for_each_question + "\n\n" + question)

        if reformatted_question is not None:
            with open(output_file, 'a', encoding='utf-8') as txt_file:
                txt_file.write(f"{i}. {question}\n")
                txt_file.write("-------------------------------\n\n")
                txt_file.write(f"{reformatted_question}\n\n")
                txt_file.write("*************************************")

        else:
            print(f"Error while reformatting question: {str(e)}")


# def convert_docx_to_txt(input_file, output_file):

#     if output_file.lower().endswith(".docx"):
#         output_file = output_file.replace("docx", "txt")

#     print("converting docx to txt \n" + input_file)
#     # check if docx file
#     if input_file.lower().endswith(".docx"):
#         try:
#             output = pypandoc.convert_file(
#                 input_file, 'plain', outputfile=output_file, encoding="utf-8")
#         except Exception as e:

#             text = docx2txt.process(input_file)
#             text = re.sub(r'(^|\n)\t', r'\1', text)
#             text = re.sub(r'(^|\n) ', r'\1', text)
#             # text = text.replace('\t', ' ')
#             text = text.strip()

#             with open(output_file, "w") as text_file:
#                 text_file.write(text)

#     return 0



def convert_docx_to_txt(input_file, output_file):

    if output_file.lower().endswith(".docx"):
        output_file = output_file.replace("docx", "txt")

    print("converting docx to txt \n" + input_file)
    
    # check if docx file
    if input_file.lower().endswith(".docx"):
        text = ""
        
        # try pypandoc conversion first
        try:
            pypandoc.convert_file(
                input_file, 'plain', outputfile=output_file, encoding="utf-8")
            with open(output_file, 'r', encoding='utf-8') as f:
                text = f.read()
        except Exception as e:
            # if pypandoc conversion fails, use docx2txt as a fallback
            text = docx2txt.process(input_file)
            text = re.sub(r'(^|\n)\t', r'\1', text)
            text = re.sub(r'(^|\n) ', r'\1', text)
            text = text.strip()

        # always remove dash lines
        with open(output_file, 'w', encoding='utf-8') as f_out:
            for line in text.split('\n'):
                if not re.match(r'^  -* -*\n', line):
                    f_out.write(line + '\n')

    return 0

def delete_dash_lines(input_file, output_file):
    with open(input_file, 'r' , encoding='utf-8') as f_in, open(output_file, 'w', encoding='utf-8') as f_out:
        for line in f_in:
            if not re.match(r'^  -* -*\n', line):
                f_out.write(line)
    return 0


def convert_txt_to_docx(input_file, output_file):

    try:
        if output_file.lower().endswith(".txt"):
            output_file = output_file.replace("txt", "docx")

        print("converting txt to docx \n" + input_file)
        # # check if txt file
        if input_file.lower().endswith(".txt"):

            with open(input_file, 'r', encoding='utf-8') as file:
                filedata = file.read()
            filedata = filedata.replace('Exp: ------', 'Exp: \n------')
            lines = filedata.splitlines()
            # lines = [line for line in lines if '--------- ---------' not in line]

            doc = docx.Document()

            table_started = False
            table_rows = []

            for line in lines:
                line = line.strip()
                print(line)
                if "---------" in line:
                    if table_started:
                        # End of table
                        table = doc.add_table(
                            rows=len(table_rows), cols=len(table_rows[0]))
                        for i, row in enumerate(table_rows):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = cell
                        table_started = False
                        table_rows = []
                    else:
                        # Start of table
                        table_started = True
                elif table_started:
                    # Table row
                    # Split on 2 or more whitespace characters
                    row = re.split(r'\s{2,}', line)
                    table_rows.append(row)
                else:
                    # Normal text
                    doc.add_paragraph(line)

            if table_started:
                # End of last table
                table = doc.add_table(rows=len(table_rows), cols=2)
                for i, row in enumerate(table_rows):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell

                # Define the desired font style and size
            font_name = 'Sahitya'
            font_size = 12
        # Change the font for all paragraphs in the document
            for paragraph in doc.paragraphs:
                set_font(paragraph, font_name, font_size)

                # Change the font for all cells in all tables in the document
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            set_font(paragraph, font_name, font_size)
                    # Format all tables in the document

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        set_cell_border(cell)

    # Format all tables in the document
            for table in doc.tables:
                set_column_widths(table)

            doc.save(output_file)
        return 0

    except Exception as e:
        writeLog(f"Error while converting {input_file} txt to docx: {str(e)}")
        print(f"Error while converting {input_file} txt to docx: {str(e)}")


def set_font(paragraph, font_name, font_size):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        # Ensure the font name is set correctly
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)


def convert_docx_to_txt_pypandoc(input_file, output_file):
    # pypandoc.download_pandoc()
    text = pypandoc.convert_file(input_file, 'latex')
    with open(output_file, "w") as text_file:
        text_file.write(text)


def convert_docx_to_txt_python_docx(input_file, output_file):
    doc = docx.Document(input_file)
    text = '\n'.join([p.text for p in doc.paragraphs])
    with open(output_file, "w") as text_file:
        text_file.write(text)


def backend_insert_questions_separator(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    modified_lines = []
    for line in lines:
        if line.strip():  # Check if the line is not empty after stripping whitespace
            # Check if the line starts with the number format
            if re.match(r'^\d+\.\)', line.strip()):
                modified_lines.append("*****\n")  # Insert separator line
            modified_lines.append(line)  # Append the original line

    # Write the modified lines to a new file
    with open(output_file, 'w', encoding='utf-8') as new_file:
        new_file.writelines(modified_lines)
        new_file.write("\n*****\n")
    return 0


def convert_ratta_to_json(filename):
    writeLog(f"{filename} \n")
    full_path = filename
    filename = os.path.basename(full_path)
    if filename.endswith('.docx'):
        # Convert .docx to .txt
        convert_docx_to_txt(full_path, filename.replace('.docx', '.txt'))
        # Update input filename for further processing
        filename = filename.replace('.docx', '.txt')
        updated_file = "separated_" + filename
        backend_insert_questions_separator(filename, updated_file)
    else:
        updated_file = "separated_" + filename
        backend_insert_questions_separator(full_path, updated_file)

    # Read the content of the file
    with open(updated_file, 'r', encoding='utf-8') as file:
        text = file.read()

    # Regular expression pattern
    # pattern = r'\*{5}([\s\S]*?)\*{5}'
    pattern = r'\*{5}(.*?)(?=\*{5})'

    # Find all matches in the text
    matches = re.findall(pattern, text, re.DOTALL)

    with open(filename.replace('.txt', '.json'), 'w', encoding='utf-8') as file:
        json.dump([backend_process_question_to_json(match)
                  for match in matches], file, indent=2, ensure_ascii=False)

    if os.path.exists(updated_file):
        os.remove(updated_file)

    with open(filename.replace('.txt', '.json'), 'r', encoding='utf-8') as file:
        data = json.load(file)

    if os.path.exists(filename.replace('.txt', '.json')):
        os.remove(filename.replace('.txt', '.json'))

    if os.path.exists(filename):
        os.remove(filename)
    return data


def convert_separated_file_to_json(input_file, output_file):

    writeLog(f"{input_file} \n")
    # Read the content of the file
    with open(input_file, 'r', encoding='utf-8') as file:
        text = file.read()

    # Get the base name without extension
    output_file = os.path.splitext(output_file)[0] + ".json"

    # Regular expression pattern
    # pattern = r'\*{5}([\s\S]*?)\*{5}'
    pattern = r'\*{5}(.*?)(?=\*{5})'

    # Find all matches in the text
    matches = re.findall(pattern, text, re.DOTALL)
    # writeLog(input_file)
    with open(output_file, 'w', encoding='utf-8') as file:
        json.dump([backend_process_question_to_json(match)
                  for match in matches], file, indent=2, ensure_ascii=False)

    # with open(filename.replace('.txt', '.json'), 'r', encoding='utf-8') as file:
    #     data = json.load(file)
    #     return data

    return len(matches)


def backend_delete_after(input_string, pattern):
    # Search for the pattern "Exp:" from the end of the input string
    exp_index = input_string.rfind(pattern)

    # If "Exp:" is found, delete everything after it
    if exp_index != -1:
        non_deleted_part = input_string[:exp_index]
        deleted_part = input_string[exp_index + len(pattern):]
        return (non_deleted_part, deleted_part)

    # If "Exp:" is not found, return the input string as it is
    return (input_string, "")


def backend_process_question_to_json(question_match):
    returned = backend_delete_after(question_match, "\nExp:")
    explanation = returned[1].strip()
    returned = backend_delete_after(returned[0], "\nAns.")
    answer = returned[1].strip()
    returned = backend_delete_after(returned[0], "\n(d)")
    option_d = returned[1].strip()
    returned = backend_delete_after(returned[0], "\n(c)")
    option_c = returned[1].strip()
    returned = backend_delete_after(returned[0], "\n(b)")
    option_b = returned[1].strip()
    returned = backend_delete_after(returned[0], "\n(a)")
    option_a = returned[1].strip()

    question_content = returned[0].strip()

    answer_content = None
    if answer == "a":
        answer_content = option_a
    elif answer == "b":
        answer_content = option_b
    elif answer == "c":
        answer_content = option_c
    elif answer == "d":
        answer_content = option_d
    else:
        writeLog(question_match)

    if (question_content == "") or (option_a == "") or (option_b == "") or (option_c == "") or (option_d == "") or (answer == "") or (answer_content == None) or (answer_content == ""):
        # print(question_match)
        writeLog(question_match)

        # corrected_question = microBrain.ask_question(question_match, model="llama3-70b-8192", host="groq" , json_format= True)
        # corrected_json = json.loads(corrected_question)
        # writeLog("Corrected Version:")
        # writeLog( f"\n{corrected_json['question']}\n(a) {corrected_json['option_a']}\n(b) {corrected_json['option_b']}\n(c) {corrected_json['option_c']}\n(d) {corrected_json['option_d']}\nAns. {corrected_json['answer']}\nExp: {corrected_json['explanation']}\n")
        # time.sleep(10)

        # writeLog(corrected_question)

    return {
        'question': question_content,
        'option_a': option_a,
        'option_b': option_b,
        'option_c': option_c,
        'option_d': option_d,
        'answer': answer,
        'answer_content': answer_content,
        'explanation': explanation
    }


def convert_ratta_to_one_line(input_filename, output_filename, include_options=False):

    base_name = os.path.basename(input_filename)

    convert_ratta_to_json(input_filename)
    if base_name.endswith('.docx'):
        base_name = base_name.replace('.docx', '.json')
    else:
        # Update input filename for further processing
        base_name = base_name.replace('.txt', '.json')

    with open(base_name, 'r', newline='', encoding='utf-8') as json_file:
        data = json.load(json_file)

    # Extract the required fields and write them to a CSV file
    with open(output_filename, 'w', newline='', encoding='utf-8') as output_file:
        writer = csv.writer(output_file)

        # Write data to the CSV file
        if include_options is False:
            for item in data:
                writer.writerow([item['question'], item['answer_content']])
        else:
            for item in data:
                writer.writerow([item['question'], item['option_a'], item['option_b'], item['option_c'],
                                item['option_d'], item['answer'], item['answer_content'], item['explanation']])


def details_extracted(input_file):
    data = convert_ratta_to_json(input_file)
    print(input_file)
    print(len(data))
    length_of_question = 0
    length_of_option = 0
    length_of_explanation = 0
    number_of_question = len(data)
    number_of_question_not_having_explanation = 0
    for items in data:
        length_of_question += len(items['question'].split())
        length_of_option += len(items['option_a'].split())
        length_of_option += len(items['option_b'].split())
        length_of_option += len(items['option_c'].split())
        length_of_option += len(items['option_d'].split())
        length_of_explanation += len(items['explanation'].split())
        if items['explanation'] == "":
            number_of_question_not_having_explanation += 1

    if number_of_question > 0:
        average_of_length_of_question = length_of_question/number_of_question
        average_of_length_of_options = length_of_option/number_of_question/4
        percentage_of_question_not_having_explanation = (
            number_of_question_not_having_explanation/number_of_question)*100
        average_of_length_of_explanation = length_of_explanation / \
            (number_of_question - number_of_question_not_having_explanation)
    return {
        "Name of File": os.path.basename(input_file),
        "Total Questions": number_of_question,
        "Percent NO Explanation": percentage_of_question_not_having_explanation,
        "No Explanation": number_of_question_not_having_explanation,
        "Question length": average_of_length_of_question,
        "Option length": average_of_length_of_options,
        "Explanation length": average_of_length_of_explanation
    }


def process_files_in_folder(folder_path, file_info_function, output_file_path):

    folder_name = os.path.basename(folder_path)
    total_questions = 0

    # Open the output file in write mode
    with open("Folder_details"+folder_name, 'w') as output_file:

        output_file.write("Folder: " + folder_name + "\n\n\n")
        output_file.write("Location of Folder: " + folder_path + "\n\n\n")
        # Iterate over files in the folder
        for filename in os.listdir(folder_path):
            # Get the full path of the file
            file_path = os.path.join(folder_path, filename)

            # Check if the path is a file (not a directory)
            if os.path.isfile(file_path):
                # Get information about the file using the provided function
                file_info = file_info_function(file_path)
                total_questions += file_info['Total Questions']

                # Write the file information to the output file
                # output_file.write(f"File: {filename}\n")
                for item in file_info:
                    output_file.write(f"{item}: {file_info[item]}\n")
                # output_file.write(file_info.__str__() )
                # Write file info followed by newline
                output_file.write("\n\n\n\n")
        output_file.write("Total questions in Folder : " +
                          str(total_questions))


def sort_options(question):
    options = [question['option_a'], question['option_b'],
               question['option_c'], question['option_d']]
    # options.sort()
    return options


def append_to_json_file(existing_file, new_entries_list):

    is_new_entry = False

    directory = os.path.dirname(existing_file)

    if not os.path.exists(directory) and directory != '':
        os.makedirs(directory)

    if not os.path.exists(existing_file):
        with open(existing_file, 'w') as f:
            f.write("[]")
            # print(f"File '{existing_file}' created successfully.")
    try:
        with open(existing_file, 'r+', encoding='utf-8') as file:
            # Load existing data
            existing_data = json.load(file)

            # Append new entries
            for entry in new_entries_list:
                if entry not in existing_data:
                    existing_data.append(entry)
                    is_new_entry = True

            # Set file position to the beginning
            file.seek(0)

            # Write the updated data back to the file
            json.dump(existing_data, file, ensure_ascii=False, indent=4)
            file.truncate()  # Truncate any remaining content (in case the new data is smaller)

        # print(f"{len(new_entries_list)} entries appended to '{existing_file}'.")
    except FileNotFoundError:
        # If the file doesn't exist, create it and write new entries to it
        with open(existing_file, 'w', encoding='utf-8') as file:
            json.dump(new_entries_list, file, ensure_ascii=False, indent=4)
        # print(f"{len(new_entries_list)} entries written to '{existing_file}'.")
    finally:
        return is_new_entry


def are_arrays_similar(arr1, arr2):

    # Check if any three consecutive elements are the same in both arrays
    for i in range(len(arr1) - 2):
        if arr1[i:i+3] == arr2[i:i+3]:
            return True
    return False


def find_D_questions_in_json_files_save_to_different_folders(file1, file2):
    """
    Finds and saves duplicate questions between two JSON files.

    Args:
        file1 (str): The path to the first JSON file.
        file2 (str): The path to the second JSON file.

    Returns:
        int: The number of duplicate questions found and saved.

    Raises:
        FileNotFoundError: If either of the input files does not exist.

    Notes:
        - The function creates new folders if they do not exist.
        - The function saves the duplicate questions in two separate folders.
        - The function saves all duplicated questions in a separate file.

    """

    # get base name of file1 and file2
    base_name1 = os.path.basename(file1)
    base_name2 = os.path.basename(file2)

    # new path for file1 and file2 in Duplicate_Folder
    duplicate_file1 = os.path.join(
        "Duplicate", 'Duplicate_Folder_1', base_name1)
    duplicate_file2 = os.path.join(
        "Duplicate", 'Duplicate_Folder_2', base_name2)
    duplicated_file = os.path.join("Duplicate", "Duplicate_Questions")

    # create new folders if not exist
    if not os.path.exists('Duplicate/Duplicate_Folder_1'):
        os.makedirs('Duplicate/Duplicate_Folder_1')

    if not os.path.exists('Duplicate/Duplicate_Folder_2'):
        os.makedirs('Duplicate/Duplicate_Folder_2')

    number_of_duplicate = 0
    with open(file1, 'r', encoding='utf-8') as f1, open(file2, 'r', encoding='utf-8') as f2:
        data1 = json.load(f1)
        data2 = json.load(f2)

        duplicate_questions_file_1 = []
        duplicate_questions_file_2 = []
        duplicated_questions = []

        for q1 in data1:
            for q2 in data2:
                if compare_two_questions(q1, q2):
                   # check length of explanation  which ever is long keep it and add other to duplicate
                    if len(q1['explanation']) >= len(q2['explanation']):
                        duplicate_questions_file_2.append(q2)
                    else:
                        duplicate_questions_file_1.append(q1)

                    duplicated_questions.append(q1)
                    duplicated_questions.append(q2)

                    number_of_duplicate += 1

        if duplicate_questions_file_1:
            append_to_json_file(duplicate_file1, duplicate_questions_file_1)
            # print(f"{len(duplicate_questions_file_1)} duplicate questions found and saved to '{duplicate_file1}'.")

        if duplicate_questions_file_2:
            append_to_json_file(duplicate_file2, duplicate_questions_file_2)
            # print(f"{len(duplicate_questions_file_2)} duplicate questions found and saved to '{duplicate_file2}'.")

        append_to_json_file(duplicated_file, duplicated_questions)
        print(
            f"{number_of_duplicate} duplicate questions found and saved to '{duplicated_file}'.")
        return number_of_duplicate


def compare_two_questions(q1, q2):
    is_duplicate = False
    if q1['option_a'] != "" and q1['option_b'] != "" and q1['option_c'] != "" and q1['option_d'] != "" and q2['option_a'] != "" and q2['option_b'] != "" and q2['option_c'] != "" and q2['option_d'] != "":
        # if q1['option_a'] is not "" and q1['option_b'] is not "" and q1['option_c'] is not "" and q1['option_d'] is not "" and q2['option_a'] is not "" and q2['option_b'] is not "" and q2['option_c'] is not "" and q2['option_d'] is not "" :
        # processed_q1_options = process_options_for_similarity(q1)
        # processed_q2_options = process_options_for_similarity(q2)

        # q1_question = re.sub(r'[^\u0900-\u097F]+', '', q1['question'])
        # q2_question = re.sub(r'[^\u0900-\u097F]+', '', q2['question'])

        # if (processed_q1_options == processed_q2_options) and q1['answer'] == q2['answer'] :
        # if (processed_q1_options == processed_q2_options) and q1['answer'] == q2['answer'] and q1_question == q2_question:
        # if q1_question== q2_question: # and q1['answer_content'] == q2['answer_content']:
        # and q1['question']== q2['question']:
        # if q1['option_a'] == q2['option_a'] and q1['option_b'] == q2['option_b'] and q1['option_c'] == q2['option_c'] and q1['option_d'] == q2['option_d'] and q1['answer'] == q2['answer'] and q1['question'].find("सुमेल") == -1:
        # if (q1['option_a']== q2['option_a'] and q1['option_b'] == q2['option_b'] and q1['option_c'] == q2['option_c'] and q1['option_d'] == q2['option_d'] and q1['answer'] == q2['answer'] ) or (q1_question== q2_question):
        # if q1['option_a']== q2['option_a'] and q1['option_b'] == q2['option_b'] and q1['option_c'] == q2['option_c'] and q1['option_d'] == q2['option_d'] and q1['answer'] == q2['answer']  and q1['question'].find("-----------") == -1 and  q1['question'].find("सुमेल") == -1:
        # if are_arrays_similar(processed_q1_options, processed_q2_options) and q1['answer'] == q2['answer']:
        # if q1['question'] == q2['question']:
        # if q1['explanation'] == q2['explanation']:
        # if q1['question'] == q2['question'] and q1['explanation'] == q2['explanation']:

        # if q1['option_a'] == q2['option_a'] and q1['option_b'] == q2['option_b'] and q1['option_c'] == q2['option_c'] and q1['option_d'] == q2['option_d'] and q1['question'] == q2['question']:
        if q1['option_a'] == q2['option_a'] and q1['option_b'] == q2['option_b'] and q1['option_c'] == q2['option_c'] and q1['option_d'] == q2['option_d'] and q1['answer'] == q2['answer']:
            is_duplicate = True
            writeLog(questionString(q1))
            writeLog(questionString(q2))
            # similarity = jaccard_similarity(q1_question, q2_question)
            # if similarity < 1.0:
            #     print(similarity)
            #     print(q1_question)
            #     print(q2_question)

    return is_duplicate


def is_question_to_be_deleted(question):
    return False
    # return true if question contains word "कथन" or  "कथनों"
    # if len(question["question"].split()) > 40:
    # if not (('\n' in question['question']) and ("सुमेल" not in question['question'] or "सूची-" not in question['question'])):
    # #     if ("कथन" in question['question'] or "कथनों" in question['question']) and question['question'].count('\n') >= 2:
    if "\n" in question['question']:  # delete if any there is a new line character
        # if   not ("सुमेल"  in question['question'] and  "सूची-"  in question['question']  ) :
        if not ("सुमेलित क" in question['question'] or "सुमेलित न" in question['question'] or question['question'].count('सूची') >= 2):
            return True
    return False


def process_options_for_similarity(question):
    options = [question['option_a'], question['option_b'],
               question['option_c'], question['option_d']]
    options = [re.sub(r'[^a-zA-Z\u0900-\u097F\d]+', '', option)
               for option in options]
    # options.sort()
    return options


def jaccard_similarity(string1, string2):
    set1 = set(string1.split())
    set2 = set(string2.split())
    intersection = len(set1.intersection(set2))
    union = len(set1.union(set2))
    similarity = intersection / union if union > 0 else 0
    return similarity


def combine_json_files(folder_path, output_file):
    combined_data = []

    # Iterate through each file in the folder
    for file_name in os.listdir(folder_path):
        if file_name.endswith('.json'):  # Check if the file is a JSON file
            file_path = os.path.join(folder_path, file_name)
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                combined_data.extend(data)  # Combine the data from all files

    # Write the combined data into a new JSON file
    with open(output_file, 'w', encoding='utf-8') as out_file:
        json.dump(combined_data, out_file, ensure_ascii=False, indent=4)

    print(f"Combined JSON data saved to '{output_file}'")


def find_duplicates_in_docx_folders_save_to_different_folders(folder1, folder2):

    temp1 = "temp1"
    temp2 = "temp2"
    temp12 = "temp12"
    temp22 = "temp22"
    process_folder_for_given_function(
        folder1, temp1, convert_docx_to_txt, "files")
    process_folder_for_given_function(
        temp1, temp1, backend_insert_questions_separator, "files")
    process_folder_for_given_function(
        temp1, temp12, convert_separated_file_to_json, "files")

    process_folder_for_given_function(
        folder2, temp2, convert_docx_to_txt, "files")
    process_folder_for_given_function(
        temp2, temp2, backend_insert_questions_separator, "files")
    process_folder_for_given_function(
        temp2, temp22, convert_separated_file_to_json, "files")

    files1 = []
    for root1, _, filenames1 in os.walk(temp12):
        for filename1 in filenames1:
            files1.append(os.path.join(root1, filename1))

    # Collect all files from folder2
    files2 = []
    for root2, _, filenames2 in os.walk(temp22):
        for filename2 in filenames2:
            files2.append(os.path.join(root2, filename2))

    # Compare files between the two lists
    duplicates = 0
    for file1_path in files1:
        for file2_path in files2:
            duplicates += find_D_questions_in_json_files_save_to_different_folders(
                file1_path, file2_path)
    print(f"{duplicates} duplicate questions found.")

    if os.path.exists(temp1):
        shutil.rmtree(temp1)
    if os.path.exists(temp2):
        shutil.rmtree(temp2)
    if os.path.exists(temp12):
        shutil.rmtree(temp12)
    if os.path.exists(temp22):
        shutil.rmtree(temp22)


def process_folder_for_given_function(input_folder, output_folder, process_function, file_or_folder="files"):
    """
    Recursively processes files in the input folder, maintains folder structure, and applies the
    specified process function to each file, saving the processed files in the output folder.
    """

    total = 0
    # Create the output folder if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # Traverse the input folder recursively
    for root, directories, files in os.walk(input_folder):
        # Construct the corresponding output subfolder
        relative_path = os.path.relpath(root, input_folder)
        output_subfolder = os.path.join(output_folder, relative_path)

        # Create the output subfolder if it doesn't exist
        if not os.path.exists(output_subfolder):
            os.makedirs(output_subfolder)
        print(process_function.__name__.__str__() + "           " + root)
        if file_or_folder == "files":
            # Process each file in the current folder
            for filename in files:
                input_file_path = os.path.join(root, filename)
                output_file_path = os.path.join(output_subfolder, filename)
                total += process_function(input_file_path, output_file_path)
                print(process_function.__name__.__str__() +
                      "           " + filename)
        else:
            total += process_function(root, output_subfolder)

    return total


def convert_json_to_txt(json_file, text_file):

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except ValueError:
        return 0

    # Get the base name without extension
    text_file = os.path.splitext(text_file)[0] + ".txt"

    with open(text_file, 'w', encoding='utf-8') as f:
        for i, question in enumerate(data, start=1):
            f.write(f"{question['question']}\n")
            f.write(f"(a) {question['option_a']}\n")
            f.write(f"(b) {question['option_b']}\n")
            f.write(f"(c) {question['option_c']}\n")
            f.write(f"(d) {question['option_d']}\n")
            f.write(f"Ans. {question['answer']}\n")
            f.write(f"Exp: {question['explanation']}\n")
            f.write("\n\n\n\n")
    return 0


def find_duplicates_in_a_single_json_folder(input_folder, output_folder):
    """
    Finds and saves duplicate questions within JSON files in a folder.

    Args:
        folder_name (str): The path to the folder containing JSON files.
    """
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    json_files = [f for f in os.listdir(input_folder) if f.endswith('.json')]

    duplicates_found = 0

    for i, file1 in enumerate(json_files):
        with open(os.path.join(input_folder, file1), 'r', encoding='utf-8') as f1:
            data1 = json.load(f1)
            duplicate_file_path1 = os.path.join(
                output_folder, os.path.splitext(file1)[0] + "_D.json")

            for file2 in json_files[i+1:]:
                with open(os.path.join(input_folder, file2), 'r', encoding='utf-8') as f2:
                    data2 = json.load(f2)
                    duplicate_file_path2 = os.path.join(
                        output_folder, os.path.splitext(file2)[0] + "_D.json")

                    for q1 in data1:

                        if is_question_to_be_deleted(q1):
                            append_to_json_file(duplicate_file_path1, [q1])
                            continue

                        for q2 in data2:

                            if is_question_to_be_deleted(q2):
                                append_to_json_file(duplicate_file_path2, [q2])
                                continue

                            if compare_two_questions(q1, q2):
                                duplicates_found += 1
                                if len(q1['explanation']) < len(q2['explanation']):
                                    append_to_json_file(
                                        duplicate_file_path1, [q1])
                                else:
                                    append_to_json_file(
                                        duplicate_file_path2, [q2])

    return duplicates_found


def save_duplicate_in_same_docx_folder(input_folder, output_folder):
    temp_folder = "temporary"
    temp_folder_2 = "temporary2"
    temp_folder_3 = "temporary3"

    try:
        # Step 1: Convert .docx files to .txt files
        process_folder_for_given_function(
            input_folder, temp_folder, convert_docx_to_txt, "files")

        # Step 2: Backend insert questions separator
        process_folder_for_given_function(
            temp_folder, temp_folder, backend_insert_questions_separator, "files")

        # Step 3: Convert separated files to JSON
        process_folder_for_given_function(
            temp_folder, temp_folder_2, convert_separated_file_to_json, "files")

        # Step 4: Find and save duplicate questions
        duplicates_found = find_Ds_in_a_single_json_folder(
            temp_folder_2, temp_folder_3)

        process_folder_for_given_function(
            temp_folder_3, output_folder, convert_json_to_txt, "files")

    except Exception as e:
        print(f"An error occurred: {e}")

    finally:
        # Delete temporary folders
        if os.path.exists(temp_folder):
            shutil.rmtree(temp_folder)
        if os.path.exists(temp_folder_2):
            shutil.rmtree(temp_folder_2)
        if os.path.exists(temp_folder_3):
            shutil.rmtree(temp_folder_3)
    return duplicates_found


def extract_stats_from_file(input_file):
    with open(input_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    length_of_question = 0
    length_of_option = 0
    length_of_explanation = 0
    number_of_question = len(data)
    number_of_question_not_having_explanation = 0

    for item in data:
        # print(item['question'])
        length_of_question += len(item['question'].split())
        length_of_option += sum(len(item[option].split())
                                for option in ['option_a', 'option_b', 'option_c', 'option_d'])
        length_of_explanation += len(item['explanation'].split())
        if not item['explanation']:
            number_of_question_not_having_explanation += 1

    if number_of_question > 0:
        average_of_length_of_question = round(
            length_of_question / number_of_question, 1)
        average_of_length_of_options = round(
            length_of_option / (number_of_question * 4), 1)
        percentage_of_question_not_having_explanation = round(
            (number_of_question_not_having_explanation / number_of_question) * 100, 1)
        if (number_of_question - number_of_question_not_having_explanation) > 0:
            average_of_length_of_explanation = round(
                length_of_explanation / (number_of_question - number_of_question_not_having_explanation), 1)
        else:
            # if all questions have an explanation, then the length of explanations is zero.from django import forms
            average_of_length_of_explanation = 0
    else:
        average_of_length_of_question = 0
        average_of_length_of_options = 0
        percentage_of_question_not_having_explanation = 0
        average_of_length_of_explanation = 0

    return {
        # "Name of File": os.path.basename(input_file),
        "Total Questions": number_of_question,
        "No Explanation": number_of_question_not_having_explanation,
        "Percent NO Explanation": percentage_of_question_not_having_explanation,
        "Question length": average_of_length_of_question,
        "Option length": average_of_length_of_options,
        "Explanation length": average_of_length_of_explanation
    }


def aggregate_stats(old_stats, new_stats):
    total_questions = old_stats["Total Questions"] + \
        new_stats["Total Questions"]
    total_no_explanation = old_stats["No Explanation"] + \
        new_stats["No Explanation"]
    if total_questions > 0:
        percentage_no_explanation = round(
            (total_no_explanation / total_questions) * 100, 1)
    else:
        percentage_no_explanation = 0

    if total_questions - total_no_explanation > 0:
        explanation_length = round((old_stats["Explanation length"] * (old_stats["Total Questions"] - old_stats["No Explanation"]) +
                                    new_stats["Explanation length"] * (new_stats["Total Questions"] - new_stats["No Explanation"])) /
                                   (total_questions - total_no_explanation), 1)
    else:
        explanation_length = 0

    return {
        "Total Questions": total_questions,
        "No Explanation": total_no_explanation,
        "Question length": round((old_stats["Question length"] * old_stats["Total Questions"] +
                                  new_stats["Total Questions"] * new_stats["Question length"]) / total_questions, 1) if total_questions > 0 else 0,
        "Percent NO Explanation": percentage_no_explanation,
        "Option length": round((old_stats["Option length"] * old_stats["Total Questions"] +
                                new_stats["Total Questions"] * new_stats["Option length"]) / total_questions, 1) if total_questions > 0 else 0,
        "Explanation length": explanation_length
    }


def extract_stats_from_folder(folder_path):
    folder_data = {"Folder": os.path.basename(folder_path)}
    if os.path.isdir(folder_path):
        folder_stats = {
            "Total Questions": 0,
            "Percent NO Explanation": 0,
            "No Explanation": 0,
            "Question length": 0,
            "Option length": 0,
            "Explanation length": 0
        }
        children = []
        for item in sorted(os.listdir(folder_path)):
            item_path = os.path.join(folder_path, item)
            if os.path.isdir(item_path):
                child_data, child_stats = extract_stats_from_folder(item_path)
                children.append(child_data)
                folder_stats = aggregate_stats(folder_stats, child_stats)
            elif item.endswith('.json'):
                stats = extract_stats_from_file(item_path)
                children.append({"File": item, **stats})
                folder_stats = aggregate_stats(folder_stats, stats)

        folder_data.update(folder_stats)
        folder_data["Children"] = children
    return folder_data, folder_stats


def write_json(data, output_file):
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


def write_yaml(data, output_file):
    with open(output_file, 'w', encoding='utf-8') as f:
        yaml.dump(data, f, allow_unicode=True, indent=8,
                  default_flow_style=False, sort_keys=False)
        f.write('\n')


def final_function_for_deduplication(folder_path):

    temp_folder = "temporary"
    process_folder_for_given_function(
        folder_path, temp_folder, convert_docx_to_txt, "files")

    process_folder_for_given_function(
        temp_folder, temp_folder, backend_insert_questions_separator, "files")

    process_folder_for_given_function(
        temp_folder, folder_path, convert_separated_file_to_json, "files")

    json_files = glob.glob(os.path.join(
        folder_path, '**/*.json'), recursive=True)

    duplicates_found = 0
    questions_to_be_deleted_found = 0

    for j, file1 in enumerate(json_files):

        duplicate_question_in_file1 = 0
        question_to_be_deleted_in_file1 = 0

        with open(os.path.join(folder_path, file1), 'r', encoding='utf-8') as f1:
            data1 = json.load(f1)
            duplicate_file_path1 = os.path.join(
                folder_path, os.path.splitext(file1)[0] + "_D.json")
            print(file1)

            for file2 in json_files[j:]:

                print(f"        " + str(os.path.basename(file2)))

                if file1 == file2:
                    for i,  q1 in enumerate(data1):
                        if is_question_to_be_deleted(q1):
                            if append_to_json_file(duplicate_file_path1, [q1]):
                                questions_to_be_deleted_found += 1
                                question_to_be_deleted_in_file1 += 1
                                # Logging the process.
                                # writeLog(f"To be Deleted:")
                                # writeLog(
                                # f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                            continue
                        for q2 in data1[i+1:]:
                            if compare_two_questions(q1, q2):

                                # Logging the process.
                                # writeLog(f"Found Duplicate in Same File: ")
                                # writeLog(
                                #     f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                # writeLog(
                                #     f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")

                                if len(q1['explanation']) < len(q2['explanation']):
                                    if append_to_json_file(duplicate_file_path1, [q1]):
                                        duplicates_found += 1
                                        duplicate_question_in_file1 += 1
                                        # writeLog(  f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                else:
                                    if append_to_json_file(duplicate_file_path1, [q2]):
                                        duplicates_found += 1
                                        duplicate_question_in_file1 += 1
                                        # writeLog( f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")
                else:
                    with open(os.path.join(folder_path, file2), 'r', encoding='utf-8') as f2:
                        data2 = json.load(f2)
                        duplicate_file_path2 = os.path.join(
                            folder_path, os.path.splitext(file2)[0] + "_D.json")

                        for q1 in data1:

                            if is_question_to_be_deleted(q1):
                                if append_to_json_file(duplicate_file_path1, [q1]):
                                    questions_to_be_deleted_found += 1
                                    question_to_be_deleted_in_file1 += 1
                                    # writeLog("To be Deleted: ")
                                    # writeLog(
                                    # f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")

                                continue

                            for q2 in data2:

                                if is_question_to_be_deleted(q2):
                                    if append_to_json_file(duplicate_file_path2, [q2]):
                                        questions_to_be_deleted_found += 1
                                        # writeLog("To be Deleted: ")
                                        # writeLog(
                                        # f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")
                                    continue

                                if compare_two_questions(q1, q2):
                                    # writeLog("Duplicate Found : ")
                                    # writeLog(file1)
                                    # writeLog(
                                    #     f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                    # writeLog(file2)
                                    # writeLog(
                                    # f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")

                                    if len(q1['explanation']) < len(q2['explanation']):

                                        if append_to_json_file(duplicate_file_path1, [q1]):
                                            duplicates_found += 1
                                            duplicate_question_in_file1 += 1
                                            # writeLog(  f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n")
                                    else:

                                        if append_to_json_file(duplicate_file_path2, [q2]):
                                            duplicates_found += 1
                                            # writeLog( f"{q2['question']}\n(a) {q2['option_a']}\n(b) {q2['option_b']}\n(c) {q2['option_c']}\n(d) {q2['option_d']}\nAns. {q2['answer']}\nExp: {q2['explanation']}\n")
            # print(f"Duplicate: {duplicate_question_in_file1}")
            print(
                f"Questions to be deleted: {question_to_be_deleted_in_file1}")
    for file_path in json_files:
        os.remove(file_path)

    process_folder_for_given_function(
        folder_path, folder_path, convert_json_to_txt, "files")

    json_files = glob.glob(os.path.join(
        folder_path, '**/*.json'), recursive=True)
    for file_path in json_files:
        os.remove(file_path)
    if os.path.exists(temp_folder):
        shutil.rmtree(temp_folder)

    return duplicates_found, questions_to_be_deleted_found


def questionString(q1):
    return f"{q1['question']}\n(a) {q1['option_a']}\n(b) {q1['option_b']}\n(c) {q1['option_c']}\n(d) {q1['option_d']}\nAns. {q1['answer']}\nExp: {q1['explanation']}\n"


def get_initial_idea_of_files_from_a_folder_recursively(folder_path):

    temp_folder = "temporary"
    another_temp_folder = os.path.basename(folder_path)
    process_folder_for_given_function(
        folder_path, temp_folder, convert_docx_to_txt, "files")

    # Step 2: Backend insert questions separator
    process_folder_for_given_function(
        temp_folder, temp_folder, backend_insert_questions_separator, "files")

    # Step 3: Convert separated files to JSON
    process_folder_for_given_function(
        temp_folder, another_temp_folder, convert_separated_file_to_json, "files")

    json_output = os.path.basename(folder_path) + ".json"
    text_output = os.path.basename(folder_path) + ".text"
    folder_structure, _ = extract_stats_from_folder(another_temp_folder)
    write_json(folder_structure, json_output)
    make_json_readable(json_output, text_output)

    # delete temp folder
    if os.path.exists(temp_folder):
        shutil.rmtree(temp_folder)

    if os.path.exists(another_temp_folder):
        shutil.rmtree(another_temp_folder)


def visualize_json(json_data, parent_node=None, graph=None):
    if graph is None:
        graph = Digraph()

    if isinstance(json_data, dict):
        for key, value in json_data.items():
            if parent_node:
                graph.edge(parent_node, key)
            visualize_json(value, key, graph)
    elif isinstance(json_data, list):
        for index, item in enumerate(json_data):
            visualize_json(item, f'{parent_node}[{index}]', graph)
    else:
        graph.node(parent_node, str(json_data))

    return graph


def json_to_xml(json_data):
    def parse_dict(data, parent):
        for key, value in data.items():
            if key != "Children":
                # Replace special characters in attribute names
                formatted_key = key.replace(" ", "_").replace("%", "percent")
                parent.set(formatted_key, str(value))
            else:
                for child in value:
                    child_node = ET.Element("node")
                    parent.append(child_node)
                    parse_dict(child, child_node)

    # Create root node
    root = ET.Element("node")

    # Parse JSON data recursively
    parse_dict(json_data, root)

    # Create XML tree
    tree = ET.ElementTree(root)

    # Serialize XML to string
    xml_str = ET.tostring(root, encoding="unicode")

    return xml_str

    output_file = os.path.basename(folder_path) + ".xml"
    with open(output_file, 'w') as f:
        f.write(xml_str)


def update_file_names_in_a_folder(folder_path, words_to_remove):
    # Iterate over all files and directories in the given folder
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            # Get the full path of the file
            file_path = os.path.join(root, file)
            print(len(file_path))
            # Extract the file name and file extension
            file_name, file_extension = os.path.splitext(file)
            print(len(file_name))
            print(file_name)
            # Check if any of the words in words_to_remove are in the file name
            for word_to_remove in words_to_remove:
                if word_to_remove in file_name:
                    # Generate the new file name by removing the word_to_remove
                    file_name = file_name.replace(word_to_remove, '')
            # if len(file_name) > 80:
            #     file_name  =  file_name[:80]

            # Construct the new file path
            new_file_path = os.path.join(root, file_name + file_extension)
            try:
                # Rename the file if the new file name is different from the old one
                if file_path != new_file_path:
                    os.rename(file_path, new_file_path)
                    print(
                        f"Renamed '{file}' to '{file_name + file_extension}'")
            except Exception as e:
                print(f"Failed to rename '{file}': {e}")


def make_json_readable(json_file_path, text_file_path):
    # Read complete file as single string
    with open(json_file_path, 'r') as file:
        data = file.read()

    # Remove square brackets [] and curly braces {}
    cleaned_text1 = re.sub(r'[\[\]{}]', '', data)

    # Remove all occurrences of "Children"
    cleaned_text2 = re.sub(r'"Children":', '', cleaned_text1)

    # Remove lines containing the word "length"
    cleaned_text3 = '\n'.join(
        line for line in cleaned_text2.split('\n') if 'length' not in line)

    # Remove lines containing only a comma
    cleaned_text4 = re.sub(r'^\s*,\s*$', '', cleaned_text3, flags=re.MULTILINE)

    # Remove specific pattern: comma + newline + any number of spaces + double quote
    # cleaned_text5 = re.sub(r',\n\s*"(?!\w)', '', cleaned_text4,  flags=re.MULTILINE)
    # cleaned_text5 = re.sub(r',\s*"\w+":', ', ', cleaned_text4)
    cleaned_text5 = re.sub(r'\n\s*"Total Questions":', ' "T":', cleaned_text4)
    cleaned_text6 = re.sub(r'\n\s*"No Explanation":', ' "N":', cleaned_text5)
    cleaned_text7 = re.sub(
        r'\n\s*"Percent NO Explanation":', ' "P":', cleaned_text6)
    cleaned_text8 = re.sub(r'\n\s*\n', '\n', cleaned_text7)
    cleaned_text9 = re.sub(r'.json", "T":', ' - ', cleaned_text8)
    cleaned_text10 = re.sub(r'", "T":', ' = ', cleaned_text9)
    cleaned_text11 = re.sub(r'"P":', '', cleaned_text10)
    cleaned_text12 = re.sub(r'"N":', '', cleaned_text11)
    cleaned_text13 = re.sub(r'"', '', cleaned_text12)

    # Write cleaned text to text file
    with open(text_file_path, 'w') as file:
        file.write(cleaned_text13)


def segregate_question(folder_path, output_folder_path, guide_line_json, search_in_explanation=False):

    if not os.path.exists(output_folder_path):
        os.makedirs(output_folder_path)

    temp_folder = "temporary"
    another_temp = "another_temp"
    process_folder_for_given_function(
        folder_path, temp_folder, convert_docx_to_txt, "files")

    process_folder_for_given_function(
        temp_folder, temp_folder, backend_insert_questions_separator, "files")

    process_folder_for_given_function(
        temp_folder, folder_path, convert_separated_file_to_json, "files")

    json_files = glob.glob(os.path.join(
        folder_path, '**/*.json'), recursive=True)

    for filename in json_files:
        basename = os.path.basename(filename)
        with open(os.path.join(folder_path, filename), 'r') as f:
            data = json.load(f)
            for question in data:
                print(questionString(question))
                # writeLog(f"\n{question['question']}\n(a) {question['option_a']}\n(b) {question['option_b']}\n(c) {question['option_c']}\n(d) {question['option_d']}\nAns.{question['answer']}\nExp: {question['explanation']}\n\n")
                question_content = question['question']
                explanation = question['explanation']
                answer_content = question['answer_content']
                # is_manual = False
                writeLog(question)
                # if "---------" in question_content or  "---------" in explanation:
                #     is_manual = True
                matched = False
                for key, values in guide_line_json.items():
                    for value in values:
                        if (value in question_content) or (value in answer_content):
                            # if (value in question_content) or (value in answer_content) or ( value in explanation):
                            # if is_manual:
                            #     writeLog(f"From {basename} to {key}")
                            #     writeLog(f"\n{question['question']}\n(a) {question['option_a']}\n(b) {question['option_b']}\n(c) {question['option_c']}\n(d) {question['option_d']}\nAns.{question['answer']}\nExp: {question['explanation']}\n\n")
                            # else:
                            with open(os.path.join(output_folder_path, key + '.txt'), 'a') as f:
                                f.write(
                                    f"\n{question['question']}\n(a) {question['option_a']}\n(b) {question['option_b']}\n(c) {question['option_c']}\n(d) {question['option_d']}\nAns.{question['answer']}\nExp: {question['explanation']}\n")
                            matched = True
                            break
                    if matched:
                        break
                if not matched:
                    if search_in_explanation:
                        for key, values in guide_line_json.items():
                            for value in values:
                                # if (value in question_content) or (value in answer_content):
                                if (value in explanation):
                                    # if is_manual:
                                    #     writeLog(f"From {basename} to {key}")
                                    #     writeLog(f"\n{question['question']}\n(a) {question['option_a']}\n(b) {question['option_b']}\n(c) {question['option_c']}\n(d) {question['option_d']}\nAns.{question['answer']}\nExp: {question['explanation']}\n\n")
                                    # else:
                                    with open(os.path.join(output_folder_path, key + '.txt'), 'a') as f:
                                        f.write(
                                            f"\n{question['question']}\n(a) {question['option_a']}\n(b) {question['option_b']}\n(c) {question['option_c']}\n(d) {question['option_d']}\nAns.{question['answer']}\nExp: {question['explanation']}\n")
                                    matched = True
                                    break
                            if matched:
                                break
                if not matched:
                    # if is_manual:
                    #     writeLog(f"keep in {basename}")
                    #     writeLog(f"\n{question['question']}\n(a) {question['option_a']}\n(b) {question['option_b']}\n(c) {question['option_c']}\n(d) {question['option_d']}\nAns.{question['answer']}\nExp: {question['explanation']}\n\n")
                    # else:
                    with open(os.path.join(output_folder_path, basename.replace("json", "txt")), 'a') as f:
                        f.write(
                            f"\n{question['question']}\n(a) {question['option_a']}\n(b) {question['option_b']}\n(c) {question['option_c']}\n(d) {question['option_d']}\nAns.{question['answer']}\nExp: {question['explanation']}\n")

    temp = "jasdkjfakljalksjdg"
    # process_folder_for_given_function(output_folder_path,temp,delete_dash_lines,"files")
    process_folder_for_given_function(
        output_folder_path, output_folder_path, convert_txt_to_docx, "files")

    txt_files = glob.glob(os.path.join(
        output_folder_path, '**/*.txt'), recursive=True)
    for file_path in txt_files:
        os.remove(file_path)

    for file_path in json_files:
        os.remove(file_path)
    if os.path.exists(temp_folder):
        shutil.rmtree(temp_folder)

    if os.path.exists(temp):
        shutil.rmtree(temp)


def set_cell_border(cell):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for border in ['top', 'left', 'bottom', 'right']:
        tcBorder = OxmlElement(f'w:{border}')
        tcBorder.set(qn('w:val'), 'single')
        tcBorder.set(qn('w:sz'), '4')
        tcBorder.set(qn('w:space'), '0')
        tcBorder.set(qn('w:color'), '000000')
        tcBorders.append(tcBorder)


# Function to calculate the maximum width of text in a column
def get_max_column_width(table, col_idx):
    max_width = 0
    for row in table.rows:
        cell = row.cells[col_idx]
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                max_width = max(max_width, len(run.text))
    return max_width

# Function to set column widths based on the widest content


def set_column_widths(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    num_columns = len(table.columns)
    for col_idx in range(num_columns):
        max_width = get_max_column_width(table, col_idx)
        if max_width > 40:
            max_width = 40
        # Convert character count to inches (adjust as needed)
        max_width_inches = max_width * 0.10  # Adjust this factor as needed
        # writeLog(max_width)
        table.columns[col_idx].width = Inches(max_width_inches)


def change_font(input_file, output_file):
    try:
        # Load the document
        doc = Document(input_file)
        # writeLog(f"Loaded document: {input_file}")

        # def set_cell_border(cell):
        #     tc = cell._element.tcPr
        #     tc.append(OxmlElement('w:tcBorders'))
        #     for border in ['top', 'left', 'bottom', 'right']:
        #         tcBorders = tc[-1]
        #         tcBorder = OxmlElement(f'w:{border}')
        #         tcBorder.set(qn('w:val'), 'single')
        #         tcBorder.set(qn('w:sz'), '4')
        #         tcBorder.set(qn('w:space'), '0')
        #         tcBorder.set(qn('w:color'), '000000')
        #         tcBorders.append(tcBorder)

        # Define the desired font style and size
        font_name = 'Sahitya'
        font_size = 12

        # Change the font for all paragraphs in the document
        for paragraph in doc.paragraphs:
            set_font(paragraph, font_name, font_size)

        # writeLog(f"Applied font settings to paragraphs")
        # Change the font for all cells in all tables in the document

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        set_font(paragraph, font_name, font_size)

        # writeLog(f"Applied font settings to tables")

        # Format all tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(cell)

        # writeLog("border settings to tables")

    # Format all tables in the document
        # for table in doc.tables:
        #     set_column_widths(table)

        # writeLog("column width settings to tables")

        # Save the modified document
        doc.save(output_file)
        # writeLog(f"Saved modified document as: {output_file}")

    except Exception as e:

        writeLog(f"An error occurred: {e}")
    return 0


def delete_duplicate_question(input_file, output_file):

    try:
        if input_file.lower().endswith(".txt"):
            return 0

        txt_file = os.path.basename(input_file).replace('.docx', '_D.txt')
        txt_file = os.path.join(os.path.dirname(input_file), txt_file)
        if not os.path.exists(txt_file):
            shutil.copy(input_file, output_file)
            return 0

        original_data = convert_ratta_to_json(input_file)

        duplicate_data = convert_ratta_to_json(txt_file)

        # delete duplicate json object from original_data
        filtered_data = [
            obj for obj in original_data if obj not in duplicate_data]

        # Get the base name without extension
        temp_file = "temporary_file.txt"

        with open(temp_file, 'w', encoding='utf-8') as f:
            for i, question in enumerate(filtered_data, start=1):
                f.write(f"{question['question']}\n")
                f.write(f"(a) {question['option_a']}\n")
                f.write(f"(b) {question['option_b']}\n")
                f.write(f"(c) {question['option_c']}\n")
                f.write(f"(d) {question['option_d']}\n")
                f.write(f"Ans. {question['answer']}\n")
                f.write(f"Exp: {question['explanation']}\n")
                f.write("\n\n\n\n")

        convert_txt_to_docx(temp_file, output_file)

        os.remove(temp_file)

    except Exception as e:
        writeLog(f"An error occurred: {e}")
    return 0


def final_function_to_delete_duplicate_questions(input_folder_path, output_folder_path,):
    process_folder_for_given_function(
        input_folder_path, output_folder_path, delete_duplicate_question, "files")


def testing_table_structure(input_folder_path):
    temp1 = "temp1"
    temp2 = "temp2"
    temp3 = "temp3"

    temp4 = "temp4"
    temp5 = "temp5"

    # process_folder_for_given_function(input_folder_path, temp_folder, change_font, "files")

    process_folder_for_given_function(
        input_folder_path, temp1, convert_docx_to_txt, "files")

    process_folder_for_given_function(
        temp1, temp2, backend_insert_questions_separator, "files")

    # # Step 3: Convert separated files to JSON
    process_folder_for_given_function(
        temp2, temp3, convert_separated_file_to_json, "files")

    process_folder_for_given_function(
        temp3, temp4, convert_json_to_txt, "files")

    process_folder_for_given_function(
        temp4, temp5, convert_txt_to_docx, "files")

    # shutil.rmtree(temp_folder)
    # shutil.rmtree(another_temp)
    # shutil.rmtree(evne_another_temp)


def add_space_to_table_cells(input_file, output_file):

    font_name = 'Sahitya'
    font_size = 12
    doc = docx.Document(input_file)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text += ' '
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_font(paragraph, font_name, font_size)
    doc.save(output_file)
    return 0


# def replace_single_column_tables_with_paragraphs(input_file, output_file):
#     doc = docx.Document(input_file)
#     for table in doc.tables:
#         if len(table.columns) == 1:
#             para = doc.add_paragraph()
#             para.style.font.name = 'Sahitya'
#             para.style.font.size = docx.shared.Pt(12)
#             for row in table.rows:
#                 para.add_run(row.cells[0].text + '\n')
#             table._element.clear CONTENT(0)  # Remove the table
#     doc.save(output_file)
#     return 0


def process_tables(input_file, output_file):
    doc = docx.Document(input_file)
    for tbl in doc.tables:
        tbl_index = doc.element.body.index(tbl._element)
        if len(tbl.columns) == 1:
            para = doc.add_paragraph('')
            for row in tbl.rows:
                para.add_run(row.cells[0].text + '\n')
            doc.element.body.insert(tbl_index, para._element)
            tbl._element.getparent().remove(tbl._element)  # Remove the table
        else:
            for row in tbl.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.text = para.text + ' '

    font_name = 'Sahitya'
    font_size = 12

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_font(paragraph, font_name, font_size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell)

        # Change the font for all paragraphs in the document
    for paragraph in doc.paragraphs:
        set_font(paragraph, font_name, font_size)
    doc.save(output_file)
    return 0
